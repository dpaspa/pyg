#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Ver By               Date        CC        Note                              #
# 2   David Paspa      27-Jul-2018 NA        Split from document generator     #
#                                            into separate module. Added       #
#                                            multiple prefix handling.         #
# 1   David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import argparse
import re
from enum import Enum
from docx import Document
import os.path
import sys
import traceback
from tqdm import trange

import logging
logging.basicConfig(level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Declare the application title and calling arguments help:                    #
#------------------------------------------------------------------------------#
appTitle = 'Document Generator - Reference number renumbering'
appVersion = '1'
parser = argparse.ArgumentParser(description='Renumbers the reference numbers in a RefNum document')
parser.add_argument('-d','--document', help='Input document to renumber', required=True)
parser.add_argument('-p','--prefix', help='The document reference number prefix', required=True)
args = vars(parser.parse_args())

#------------------------------------------------------------------------------#
# Declare global program variables:                                            #
#------------------------------------------------------------------------------#
iRefNum = 1
sRefPrefix = ''

#------------------------------------------------------------------------------#
# Create a progress bar:                                                       #
#------------------------------------------------------------------------------#
ps = trange(1, desc='Initialise', leave=False)

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Declare global variables:                                                #
    #--------------------------------------------------------------------------#
    global ps

    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Close the progress bar:                                                  #
    #--------------------------------------------------------------------------#
    ps.close()

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print(appTitle + ' Version ' + appVersion + '\r\n' + 'ERROR ' +
          str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
#    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    fileNotExist                       = -1
    noPrefixFound                      = -2
    pathNotExist                       = -3

errorMessage = {
    errorCode.fileNotExist             : 'Document file @1 does not exist.',
    errorCode.noPrefixFound            : 'No reference numbers found with the specified prefix @1.',
    errorCode.pathNotExist             : 'Path @1 does not exist.'
}

#------------------------------------------------------------------------------#
# Function main                                                                #
#                                                                              #
# Description:                                                                 #
# The main entry point for the program.                                        #
#------------------------------------------------------------------------------#
def main():
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = main.__name__

    #--------------------------------------------------------------------------#
    # Use the global variables:                                                #
    #--------------------------------------------------------------------------#
    global iRefNum
    global sRefPrefix

    #--------------------------------------------------------------------------#
    # Get the calling arguments:                                               #
    #--------------------------------------------------------------------------#
    sDocument = args['document']
    sRefPrefix = args['prefix']
    sRefPrefix = sRefPrefix.split(',')

    #--------------------------------------------------------------------------#
    # Open the document object:                                                #
    #--------------------------------------------------------------------------#
    d = gDoc(sDocument)

    #--------------------------------------------------------------------------#
    # Renumber any reference numbers:                                          #
    #--------------------------------------------------------------------------#
    d.tagRenumber()

    #--------------------------------------------------------------------------#
    # Save the document:                                                       #
    #--------------------------------------------------------------------------#
    d.document.save(d.inputFile)

    #--------------------------------------------------------------------------#
    # Output a success message:                                                #
    #--------------------------------------------------------------------------#
    print('Congratulations! Reference numbers renumbered successfully.')

#------------------------------------------------------------------------------#
# Class: gDoc                                                                  #
#                                                                              #
# Description:                                                                 #
# Opens the specified document as an object.                                   #
#------------------------------------------------------------------------------#
# Calling Attributes:                                                          #
# inputFile             The input template file name to use for the document.  #
#------------------------------------------------------------------------------#
class gDoc(object):
    #--------------------------------------------------------------------------#
    # Constructor:                                                             #
    #--------------------------------------------------------------------------#
    def __init__(self, inputFile):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'gDoc'

        #----------------------------------------------------------------------#
        # Make sure the input document file exists:                            #
        #----------------------------------------------------------------------#
        if not os.path.exists(inputFile):
            errorHandler(self.errProc, errorCode.fileNotExist, inputFile)

        #----------------------------------------------------------------------#
        # Set the instance attributes:                                         #
        #----------------------------------------------------------------------#
        self.inputFile = inputFile

        #----------------------------------------------------------------------#
        # Open the input document:                                             #
        #----------------------------------------------------------------------#
        self.document = Document(self.inputFile)

    #--------------------------------------------------------------------------#
    # Function: tagRenumber                                                    #
    #                                                                          #
    # Description:                                                             #
    # Renumbers the tag list.                                                  #
    #--------------------------------------------------------------------------#
    def tagRenumber(self):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'tagRenumber'

        #----------------------------------------------------------------------#
        # Use global tag number variable:                                      #
        #----------------------------------------------------------------------#
        global ps
        global iRefNum
        global sRefPrefix

        #----------------------------------------------------------------------#
        # Get the progress weighting:                                          #
        #----------------------------------------------------------------------#
        bFoundPrefix = False
        numTables = len(self.document.tables)

        #----------------------------------------------------------------------#
        # Loop through all the tables in the document table collection:        #
        #----------------------------------------------------------------------#
        iTable = 0
        for table in self.document.tables:
            iRow = 0
            iTable = iTable + 1
            for row in table.rows:
                numRows = len(table.rows)
                iRow = iRow + 1
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #------------------------------------------------------#
                        # Check if a valid tag number:                         #
                        #------------------------------------------------------#
                        s = paragraph.text
                        for p in sRefPrefix:
                            if (len(s) >= 1 + len(p) and s[:len(p)] == p):
    #                            pa = re.compile('[a-zA-Z]')
                                pn = re.compile('[0-9]')
    #                            ma = pa.match(s[:len(sRefPrefix)])
                                mn = pn.match(s[len(p):])
                                if (not mn is None):
                                    bFoundPrefix = True
                                    style = paragraph.style
                                    paragraph.text = ''
                                    sn = p + str(iRefNum)
                                    run = paragraph.add_run(sn)
                                    paragraph.style = style
                                    iRefNum = iRefNum + 1

                #--------------------------------------------------------------#
                # Update the progress bar:                                     #
                #--------------------------------------------------------------#
                pc = 1.0 / (numTables * numRows)
                ps.update(pc)
                ps.set_description('Renumbering table ' + str(iTable) + ' row ' + str(iRow))
                ps.refresh()

        #----------------------------------------------------------------------#
        # Report completion regardless of error:                               #
        #----------------------------------------------------------------------#
        if (bFoundPrefix):
            ps.set_description('Processing complete')
            ps.refresh()
            ps.close()
        else:
            errorHandler(self.errProc, errorCode.noPrefixFound, sRefPrefix)

#------------------------------------------------------------------------------#
# Call the main function:                                                      #
#------------------------------------------------------------------------------#
main()
