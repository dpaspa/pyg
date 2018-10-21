#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Ver By               Date        CC        Note                              #
# 2   David Paspa      27-Jul-2018 NA        Split from document generator     #
#                                            into separate module. Added       #
#                                            multiple prefix handling.         #
# 1   David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import re
from enum import Enum
import sys
from docx import Document
import traceback
from tqdm import trange

import logging
logging.basicConfig(filename='fdg.log',level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Declare global variables:                                                    #
#------------------------------------------------------------------------------#
ps = None

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Declare global variables:                                                #
    #--------------------------------------------------------------------------#
    global ps

    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Close the progress bar:                                                  #
    #--------------------------------------------------------------------------#
    ps.close()

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print('ERROR ' + str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
#    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    fileNotExist                       = -1
    noPrefixFound                      = -2
    pathNotExist                       = -3

errorMessage = {
    errorCode.fileNotExist             : 'Document file @1 does not exist.',
    errorCode.noPrefixFound            : 'No reference numbers found with the specified prefix @1.',
    errorCode.pathNotExist             : 'Path @1 does not exist.'
}

#------------------------------------------------------------------------------#
# Function: refRenumber                                                        #
#                                                                              #
# Description:                                                                 #
# The main entry point for the program.                                        #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# d                     The generic document object.                           #
# sPrefix               The valid prefix array to renumber.                    #
# iStart                The starting number, usually 1.                        #
#------------------------------------------------------------------------------#
def refRenumber(d, sPrefix, iStart):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = refRenumber.__name__

    #--------------------------------------------------------------------------#
    # Use global tag number variable:                                          #
    #--------------------------------------------------------------------------#
    global ps

    #--------------------------------------------------------------------------#
    # Create a progress bar:                                                   #
    #--------------------------------------------------------------------------#
    ps = trange(1, desc='Initialise', leave=False)

    #--------------------------------------------------------------------------#
    # Get the calling arguments:                                               #
    #--------------------------------------------------------------------------#
    sRefPrefix = sPrefix.split(',')

    #--------------------------------------------------------------------------#
    # Define the procedure name and trap any programming errors:               #
    #--------------------------------------------------------------------------#
    errProc = 'tagRenumber'

    #--------------------------------------------------------------------------#
    # Get the progress weighting:                                              #
    #--------------------------------------------------------------------------#
    if (iStart < 1):
        iStart = 1
    iRefNum = iStart
    bFoundPrefix = False
    numTables = len(d.document.tables)

    #--------------------------------------------------------------------------#
    # Loop through all the tables in the document table collection:            #
    #--------------------------------------------------------------------------#
    iTable = 0
    for table in d.document.tables:
        iRow = 0
        iTable = iTable + 1
        for row in table.rows:
            numRows = len(table.rows)
            iRow = iRow + 1
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    #----------------------------------------------------------#
                    # Check if a valid tag number:                             #
                    #----------------------------------------------------------#
                    s = paragraph.text
                    for p in sRefPrefix:
                        if (len(s) >= 1 + len(p) and s[:len(p)] == p):
#                            pa = re.compile('[a-zA-Z]')
                            pn = re.compile('[0-9]')
#                            ma = pa.match(s[:len(sRefPrefix)])
                            mn = pn.match(s[len(p):])
                            if (not mn is None):
                                bFoundPrefix = True
                                style = paragraph.style
                                paragraph.text = ''
                                sn = p + str(iRefNum)
                                run = paragraph.add_run(sn)
                                paragraph.style = style
                                iRefNum = iRefNum + 1

            #------------------------------------------------------------------#
            # Update the progress bar:                                         #
            #------------------------------------------------------------------#
            pc = 1.0 / (numTables * numRows)
            ps.update(pc)
            ps.set_description('Renumbering table ' + str(iTable) + ' row ' + str(iRow))
            ps.refresh()

    #--------------------------------------------------------------------------#
    # Report completion regardless of error:                                   #
    #--------------------------------------------------------------------------#
    if (bFoundPrefix):
        ps.set_description('Reference numbers renumbered successfully.')
        ps.refresh()
        ps.close()
    else:
        errorHandler(errProc, errorCode.noPrefixFound, sRefPrefix)

    #--------------------------------------------------------------------------#
    # Return the reference number reached:                                     #
    #--------------------------------------------------------------------------#
    return iRefNum
