#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Rev By               Date        CC        Note                              #
# 1.0 David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import argparse
import re
from enum import Enum
from docx import Document
from docx import shared
from docx.oxml.shared import OxmlElement, qn # Necessary Import
#from docx.enum.table import WD_ROW_HEIGHT
import openpyxl
import os.path
from shutil import copyfile
import sys
import traceback
from tqdm import trange
from time import sleep
import sqlite3
from xls2db import xls2db
import collections
import datetime
import PyPDF2
import reportlab
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
import subprocess
from contextlib import contextmanager
import cgSQL

import logging
logging.basicConfig(level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Declare the application title and calling arguments help:                    #
#------------------------------------------------------------------------------#
appTitle = 'Document Generator'
appVersion = '1'
parser = argparse.ArgumentParser(description='Generates a document from a configuration spreadsheet and set of document templates')
parser.add_argument('-c','--config', help='Configuration spreadsheet', required=True)
parser.add_argument('-d','--input', help='Input base document template file', required=True)
parser.add_argument('-e','--existingDB', help='Use existing database file', required=False)
parser.add_argument('-f','--dataField', help='Field name containing the record document template file', required=False)
parser.add_argument('-r','--dataRecord', help='Sheet field containing the record document template file', required=False)
parser.add_argument('-k','--key', help='The child document key field name', required=False)
parser.add_argument('-n','--addPageNumbers', help='Add page numbers to the output PDF file', required=False)
parser.add_argument('-o','--output', help='Output for the generated document file', required=True)
parser.add_argument('-s','--scope', help='The document scope', required=True)
parser.add_argument('-t','--type', help='The document type', required=True)
parser.add_argument('-p','--prefix', help='The document reference number prefix', required=False)
args = vars(parser.parse_args())

#------------------------------------------------------------------------------#
# Declare global program variables:                                            #
#------------------------------------------------------------------------------#
ps = 0
iNumPages = 1
iPageNum = 1
iRefNum = 1
sRefPrefix = ''

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print(appTitle + ' Version ' + appVersion + '\r\n' + 'ERROR ' +
          str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    cannotCommit                       = -1
    cannotConvertWorkbook              = -2
    cannotConnectDB                    = -3
    cannotCreateTable                  = -4
    cannotGetSQL                       = -5
    cannotQuery                        = -6
    fileChildNotExist                  = -27
    fileNotExist                       = -7
    pathNotExist                       = -8
    noChildKey                         = -9
    noDocProperties                    = -10
    noEndPlaceholder                   = -37
    noRecordDocument                   = -40
    nonASCIICharacter                  = -41
    unknownAttribute                   = -42

errorMessage = {
    errorCode.cannotCommit             : 'Cannot commit changes to sqlite database',
    errorCode.cannotConvertWorkbook    : 'Cannot convert workbook @1 from xlsx to sqlite db @2',
    errorCode.cannotConnectDB          : 'Cannot connect to sqlite database @1',
    errorCode.cannotCreateTable        : 'Cannot insert table @1 into sqlite database',
    errorCode.cannotGetSQL             : 'Cannot retrieve SQL query expression for attribute @1',
    errorCode.cannotQuery              : 'Query cannot execute using SQL expression @1 with parameters @2',
    errorCode.fileChildNotExist        : 'Child file @1 does not exist.',
    errorCode.fileNotExist             : 'Parent file @1 does not exist.',
    errorCode.pathNotExist             : 'Path @1 does not exist.',
    errorCode.noEndPlaceholder         : 'No valid END placholder in query string @1.',
    errorCode.noChildKey               : 'No child record key defined. Used for document number.',
    errorCode.noDocProperties          : 'Document has no properties in tblDocuments and tblRevisionHistory',
    errorCode.noRecordDocument         : 'No record document specified with record data.',
    errorCode.nonASCIICharacter        : 'Query expression returned a non ascii-encoded unicode string',
    errorCode.unknownAttribute         : 'Attribute @1 is unknown.',
}

#------------------------------------------------------------------------------#
# Create a progress bar:                                                       #
#------------------------------------------------------------------------------#
pm = trange(100, desc='Create document...', leave=False)

#------------------------------------------------------------------------------#
# Function main                                                                #
#                                                                              #
# Description:                                                                 #
# The main entry point for the program.                                        #
#------------------------------------------------------------------------------#
def main():
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = main.__name__

    #--------------------------------------------------------------------------#
    # Use the global prefix:                                                   #
    #--------------------------------------------------------------------------#
    global ps
    global sRefPrefix
    ps = trange(1, desc='Initialise', leave=False)

    #--------------------------------------------------------------------------#
    # Get the document scope and type:                                         #
    #--------------------------------------------------------------------------#
    sDocScope = args['scope'].upper()
    sDocType = args['type'].upper()

    #--------------------------------------------------------------------------#
    # Get the input file and output generated document file path and name:     #
    #--------------------------------------------------------------------------#
    sInput = args['input']
    sOutput = args['output']

    #--------------------------------------------------------------------------#
    # Get the reference number prefix if any specified:                        #
    #--------------------------------------------------------------------------#
    sRefPrefix = ''
    if (not args['prefix'] is None):
        sRefPrefix = args['prefix']

    #--------------------------------------------------------------------------#
    # Get the code configuration workbook name and check it exists:            #
    #--------------------------------------------------------------------------#
    wbName = args['config']
    if not os.path.exists(wbName):
        errorHandler(errProc, errorCode.filenotExist, wbName)

    #--------------------------------------------------------------------------#
    # Get the flag for using the existing DB if it exists:                     #
    #--------------------------------------------------------------------------#
    bUseExistingDB = False
    bUseExistingDB = args['existingDB']
    if (bUseExistingDB == 'True'):
        #----------------------------------------------------------------------#
        # Get the database name:                                               #
        #----------------------------------------------------------------------#
        dbName = os.path.dirname(wbName) + '/config.db'
    else:
        #--------------------------------------------------------------------------#
        # Copy the configuration worksbook to a new file so the base spreadsheet   #
        # cannot be affected:                                                      #
        #--------------------------------------------------------------------------#
        wbNameDB = os.path.dirname(wbName) + '/configdb.xlsx'
        copyfile(wbName, wbNameDB)

        #--------------------------------------------------------------------------#
        # Delete the sqlite database file if it already exists so it can be        #
        # created anew with refreshed data:                                        #
        #--------------------------------------------------------------------------#
        dbName = os.path.dirname(wbName) + '/dg.db'
        try:
            os.remove(dbName)
        except OSError:
            pass

        #--------------------------------------------------------------------------#
        # Convert the configuration workbook from xlsx to sqlite database format:  #
        #--------------------------------------------------------------------------#
        try:
            xls2db(wbNameDB, dbName)
        except:
            errorHandler(errProc, errorCode.cannotConvertWorkbook, wbNameDB, dbName)

    #--------------------------------------------------------------------------#
    # Connect to the new persistent sqlite database file:                      #
    #--------------------------------------------------------------------------#
    conn = ''
    try:
        conn = sqlite3.connect(dbName)
        conn.row_factory = sqlite3.Row
    except:
        errorHandler(errProc, errorCode.cannotConnectDB, dbName)

    #--------------------------------------------------------------------------#
    # Create the global parameters table:                                      #
    #--------------------------------------------------------------------------#
    try:
        c = conn.cursor()
        query = cgSQL.sql[cgSQL.sqlCode.tblCreateGlobalParameter]
        c.execute(query)
    except:
        errorHandler(errProc, errorCode.cannotCreateTable, query, 'no parameters')

    #--------------------------------------------------------------------------#
    # Create the base parent document. Don't know if will have children yet:   #
    #--------------------------------------------------------------------------#
    d = gDocParent(conn, sDocType, sDocScope, sInput, sOutput)
    dc = None

    #--------------------------------------------------------------------------#
    # Get the progress weighting:                                              #
    #--------------------------------------------------------------------------#
    pbChunks = 1.0
#    pc = pbwt * 1.0 / num
#    pm.update(pc)
    pm.update(100)
    pm.set_description(d.docScope + ' document complete')
    pm.refresh()

    #--------------------------------------------------------------------------#
    # Check if the document has base data for mulitple records:                #
    #--------------------------------------------------------------------------#
    if (d.parent):
        #----------------------------------------------------------------------#
        # Check if there is a record document. There should be if record data  #
        # is defined for a parent document:                                    #
        #----------------------------------------------------------------------#
        iChildNum = 1
        if (args['dataRecord'] is None and args['dataField'] is None):
            errorHandler(errProc, errorCode.noRecordDocument)

        if (args['dataRecord'] is None):
            bLookupChildName = False
            sFieldRecord = args['dataField']
        else:
            bLookupChildName = True
            sFieldRecord = args['dataRecord']

        #----------------------------------------------------------------------#
        # Get the child record key:                                            #
        #----------------------------------------------------------------------#
        if (args['key'] is None):
            errorHandler(errProc, errorCode.noChildKey)
        sFieldKey = args['key']
        sFieldKey = sFieldKey.upper()

        #----------------------------------------------------------------------#
        # Process each row in the cursor to append the new record document:    #
        #----------------------------------------------------------------------#
        for row in d.dataBase:
            #------------------------------------------------------------------#
            # Get the input file name for the child:                           #
            #------------------------------------------------------------------#
            if (bLookupChildName):
                sInputChild = os.path.dirname(d.inputFile) + '/' + str(row[sFieldRecord]) + '.docx'
            else:
                sInputChild = os.path.dirname(d.inputFile) + '/' + sFieldRecord + '.docx'

            #------------------------------------------------------------------#
            # Create the child record document. It won't need to be kept so    #
            # set a temporary output file name in the same directory as the    #
            # parent:                                                          #
            #------------------------------------------------------------------#
            dc = gDocChild(d, sInputChild, row, sFieldKey, row[sFieldKey], iChildNum)
            iChildNum = iChildNum + 1

            #------------------------------------------------------------------#
            # Update the progress bar:                                         #
            #------------------------------------------------------------------#
            pc = 0.8 / d.rowCount
            ps.update(pc)
            s = dc.fieldKey
            s = s.replace(s, str(row[s]))
            ps.set_description('Child ' + s)
            ps.refresh()

            #------------------------------------------------------------------#
            # Append the child record document PDF to the base document PDF:   #
            #------------------------------------------------------------------#
            d.appendPDF(dc)

    #--------------------------------------------------------------------------#
    # Commit the changes to the database and close the connection:             #
    #--------------------------------------------------------------------------#
    conn.commit()
    conn.close()

    #--------------------------------------------------------------------------#
    # Delete the output docx documents which are not needed:                   #
    #--------------------------------------------------------------------------#
#    d.docxDelete()

    if (not dc is None):
        dc.docxDelete()

    #--------------------------------------------------------------------------#
    # Add the page numbers to the output file if required:                     #
    #--------------------------------------------------------------------------#
    if (args['addPageNumbers'] == 'True'):
        pdfPageNumbers(d.outputPDFFileName)

    #--------------------------------------------------------------------------#
    # Report completion regardless of error:                                   #
    #--------------------------------------------------------------------------#
    ps.close()
    pm.set_description('Processing complete')
    pm.refresh()
    pm.close()

    #--------------------------------------------------------------------------#
    # Get the number of pages:                                                 #
    #--------------------------------------------------------------------------#

    #--------------------------------------------------------------------------#
    # Output a success message:                                                #
    #--------------------------------------------------------------------------#
    print('Congratulations! Document ' + d.docNumber + ' generated successfully.')

#------------------------------------------------------------------------------#
# Class: gDoc                                                                  #
#                                                                              #
# Description:                                                                 #
# Creates a generated document using sqlite data to populate a document        #
# template which has @@ placeholders in it. The document must be listed in the #
# database in a tblDocument list.                                              #
#                                                                              #
# THe template is filled with data and saved as a new docx file and is also    #
# output as a PDF file.                                                        #
#                                                                              #
# The document template must have the query definitions defined within it in   #
# tables with the query string in the first row. Different types of queries    #
# are handled as follows:                                                      #
#                                                                              #
# IMAGE                 An image will be inserted in the table cell.           #
# SQLBASE               Base data for the document, such as document number    #
#                       and title.                                             #
# SQLRECORD             A sub-document in a multi-record compound document.    #
#                       The first document will be the parent header document  #
#                       and each record will be appended as a child document.  #
# SQLROW                Multi-record table. The table must include a row of    #
#                       column headings followed by a row of cell data         #
#                       attributes using @@ data field placeholders.           #
# SQLSTATIC             A table of any cell arrangement to be populated        #
#                       without changing the table structure.                  #
#------------------------------------------------------------------------------#
# Calling Attributes:                                                          #
# conn                  The database connection object.                        #
# docType               The document type matching in the database.            #
# docScope              The data scope of the document.                        #
# inputFile             The input template file name to use for the document.  #
# outputFile            The docx file name to create after field replacement.  #
#                                                                              #
# Other Attributes created by the constructor:                                 #
# docNumber             The document number.                                   #
# docTitle              The document title.                                    #
# docVer                The current document version number.                   #
# docDate               The current document version changed date.             #
# outputDir             The output directory.                                  #
# outputBaseName        The base output file name without the path.            #
# outputFileName        The base output file name without the extension.       #
# outputPDFFileName     The output PDF file name.                              #
#------------------------------------------------------------------------------#
class gDoc(object):
    #--------------------------------------------------------------------------#
    # Constructor:                                                             #
    #--------------------------------------------------------------------------#
    def __init__(self, conn, docType, docScope, inputFile, outputFile):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'parentInit'

        #----------------------------------------------------------------------#
        # Make sure the input document file exists:                            #
        #----------------------------------------------------------------------#
        if not os.path.exists(inputFile):
            errorHandler(self.errProc, errorCode.fileNotExist, inputFile)

        #----------------------------------------------------------------------#
        # Set the instance attributes:                                         #
        #----------------------------------------------------------------------#
        self.conn = conn
        self.dataBase = None
        self.dataRow = None
        self.key = ''
        self.fieldKey = ''
        self.docType = docType.upper()
        self.docScope = docScope.upper()
        self.inputFile = inputFile
        self.outputFile = outputFile

        #----------------------------------------------------------------------#
        # Set the file characteristics:                                        #
        #----------------------------------------------------------------------#
        self.outputFileName = os.path.basename(outputFile)
        self.outputBaseName = os.path.splitext(self.outputFileName)[0]
        self.outputDir = os.path.dirname(self.outputFile)
        if not os.path.exists(self.outputDir):
            errorHandler(self.errProc, errorCode.pathNotExist, self.outputDir)

        #----------------------------------------------------------------------#
        # Create the document:                                                 #
        #----------------------------------------------------------------------#
        self.createDocument()

    #--------------------------------------------------------------------------#
    # Function: createDocument                                                 #
    #                                                                          #
    # Description:                                                             #
    # Creates the docx document and sets its properties and then loops through #
    # all the tables to populate them with data. Finally it saves the          #
    # completed document as both a .docx file and a .pdf file.                 #
    #--------------------------------------------------------------------------#
    def createDocument(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'createDocument'

        #----------------------------------------------------------------------#
        # Use global tag number variable:                                      #
        #----------------------------------------------------------------------#
        global iRefNum
        global sRefPrefix

        #----------------------------------------------------------------------#
        # Open the input document:                                             #
        #----------------------------------------------------------------------#
        self.document = Document(self.inputFile)

        #----------------------------------------------------------------------#
        # Set the document properties:                                         #
        #----------------------------------------------------------------------#
        self.setDocumentProperties()

        #----------------------------------------------------------------------#
        # Get the latest document version number:                              #
        #----------------------------------------------------------------------#
        self.getVersionNumber()

        #----------------------------------------------------------------------#
        # Assume not a parent document unless a base data placeholder is found:#
        #----------------------------------------------------------------------#
        self.parent = False

        #----------------------------------------------------------------------#
        # Prevent rows from breaking across pages:                             #
        #----------------------------------------------------------------------#
        self.preventDocumentBreak()

        #----------------------------------------------------------------------#
        # Proess all of the tables in the document:                            #
        #----------------------------------------------------------------------#
        self.processTables()

        #----------------------------------------------------------------------#
        # Renumber any reference numbers:                                      #
        #----------------------------------------------------------------------#
        if (len(sRefPrefix) > 0):
            self.tagRenumber()

        #----------------------------------------------------------------------#
        # Save the output document:                                            #
        #----------------------------------------------------------------------#
        self.document.save(self.outputFile)

        #----------------------------------------------------------------------#
        # Print the base document to PDF:                                      #
        #----------------------------------------------------------------------#
        self.printPDF()

    #--------------------------------------------------------------------------#
    # Function: processTables                                                  #
    #                                                                          #
    # Description:                                                             #
    # Processes all of the tables in the document looking.                     #
    #--------------------------------------------------------------------------#
    def processTables(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'processTables'

        #----------------------------------------------------------------------#
        # Process each table in the document:                                  #
        #----------------------------------------------------------------------#
        for table in self.document.tables:
            self.dataTable(table)
#        for section in self.document.sections:
#            header = section.header
#            footer = section.footer
#            for table in section.tables:
#                self.dataTable(table)

#            for table in header.tables:
#                self.dataTable(table)

#            for table in footer.tables:
#                self.dataTable(table)

    #--------------------------------------------------------------------------#
    # Function: preventDocumentBreak                                           #
    #--------------------------------------------------------------------------#
    def preventDocumentBreak(self):
        tags = self.document.element.xpath('//w:tr')
        rows = len(tags)
        for row in range(0,rows):
            tag = tags[row]                     # Specify which <w:r> tag you want
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)                   # Append in the new tag

    #--------------------------------------------------------------------------#
    # Function: dataTable                                                      #
    #                                                                          #
    # Description:                                                             #
    # Processes the current table looking for defined header keywords defining #
    # the content of the table.                                                #
    #--------------------------------------------------------------------------#
    def dataTable(self, table):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'dataTable'

        #------------------------------------------------------------------#
        # Get the table data source:                                       #
        #------------------------------------------------------------------#
        bDelRow = False
        self.currentTable = table
        self.currentRow = self.currentTable.rows[0]
        txtQuery = self.currentRow.cells[0].text

        #------------------------------------------------------------------#
        # Check for non-ascii characters. Can't be a keyword string in     #
        # that case:                                                       #
        #------------------------------------------------------------------#
        try:
            txtQuery.decode('ascii')
        except:
            #--------------------------------------------------------------#
            # Just ignore this table:                                      #
            #--------------------------------------------------------------#
            pass
        else:
            #--------------------------------------------------------------#
            # Check for the content type. Check if an image anchor:        #
            #--------------------------------------------------------------#
            if (txtQuery[:5] == 'IMAGE'):
                #----------------------------------------------------------#
                # Insert the image:                                        #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableInsertImage(txtQuery[6:])

            #--------------------------------------------------------------#
            # Check if base query data for the entire document:            #
            #--------------------------------------------------------------#
            elif (txtQuery[:7] == 'SQLBASE'):
                #----------------------------------------------------------#
                # Get the base data for the document:                      #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.getBaseData(txtQuery[8:])

            #--------------------------------------------------------------#
            # Check if a base query record table:                          #
            #--------------------------------------------------------------#
            elif (txtQuery[:9] == 'SQLRECORD'):
                #----------------------------------------------------------#
                # Must be a valid query. Insert the data:                  #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableBaseRecord()

            #--------------------------------------------------------------#
            # Check if an append table rows query:                         #
            #--------------------------------------------------------------#
            elif (txtQuery[:6] == 'SQLROW'):
                #----------------------------------------------------------#
                # Must be a valid query. Insert the data:                  #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableAddRows(txtQuery[7:])

            #--------------------------------------------------------------#
            # Check if a static table query:                               #
            #--------------------------------------------------------------#
            elif (txtQuery[:9] == 'SQLSTATIC'):
                #----------------------------------------------------------#
                # Must be a valid query. Insert the data:                  #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableStaticFields(txtQuery[10:])

            #--------------------------------------------------------------#
            # Check if an append table rows query:                         #
            #--------------------------------------------------------------#
            elif (txtQuery[:10] == 'SQLVERHIST'):
                #----------------------------------------------------------#
                # Must be a valid query. Insert the data:                  #
                #----------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                query = cgSQL.sql[cgSQL.sqlCode.VERHIST]
                self.tableAddRows(query)
            else:
                #----------------------------------------------------------#
                # Ignore other content:                                    #
                #----------------------------------------------------------#
                pass

    #--------------------------------------------------------------------------#
    # Function: getBaseData                                                    #
    #                                                                          #
    # Description:                                                             #
    # Gets the underlying document cursor for multi-record compound documents. #
    #--------------------------------------------------------------------------#
    def getBaseData(self, query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'getBaseData'

        #----------------------------------------------------------------------#
        # Flag this is a parent document:                                      #
        #----------------------------------------------------------------------#
        self.parent = True
        self.parentQuery = query

        #----------------------------------------------------------------------#
        # Execute the base data query and get the data:                        #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        self.dataBase = c.fetchall()

        #----------------------------------------------------------------------#
        # Iterate the data to get the row count:                               #
        #----------------------------------------------------------------------#
        self.rowCount = 0
        for row in self.dataBase:
            self.rowCount += 1

        #----------------------------------------------------------------------#
        # The cursor is one way so refresh it after iteration:                 #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        self.dataBase = c.fetchall()

    #--------------------------------------------------------------------------#
    # Function: setDocumentProperties                                          #
    #                                                                          #
    # Description:                                                             #
    # Sets the core document properties for document type, number and title.   #
    #--------------------------------------------------------------------------#
    def setDocumentProperties(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'setDocumentProperties'

        #----------------------------------------------------------------------#
        # Set the document type from the core properties:                      #
        #----------------------------------------------------------------------#
        self.document.core_properties.subject = self.docType

        #----------------------------------------------------------------------#
        # Get the document information:                                        #
        #----------------------------------------------------------------------#
        try:
            c = self.conn.cursor()
            query = cgSQL.sql[cgSQL.sqlCode.documentInfo]
            c.execute(query, (self.docType, self.docScope))
        except:
            errorHandler(self.errProc, errorCode.cannotQuery,
                         cgSQL.sqlCode.documentInfo, query, self.docType + ', ' + self.docScope)

        #----------------------------------------------------------------------#
        # Set the document number into the core properties:                    #
        #----------------------------------------------------------------------#
        data = c.fetchone()
        if (data is None):
            errorHandler(self.errProc, errorCode.noDocProperties)
        else:
            #------------------------------------------------------------------#
            # Set the document identity information into the core properties:  #
            #------------------------------------------------------------------#
            self.docNumber = data['docNumber']
            if (len(self.key) > 0):
                self.docNumber = self.docNumber + '-' + self.key

            self.docTitle = data['docTitle']
            self.document.core_properties.comments = self.docNumber
            self.document.core_properties.title = self.docTitle

    #--------------------------------------------------------------------------#
    # Function: getVersionNumber                                               #
    #                                                                          #
    # Description:                                                             #
    # Gets the core document properties for version number and date and retri  #
    # the document type.                                                       #
    #--------------------------------------------------------------------------#
    def getVersionNumber(self):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        self.errProc = 'getVersionNumber'

        #----------------------------------------------------------------------#
        # Get the version history for the document:                            #
        #----------------------------------------------------------------------#
        try:
            c = self.conn.cursor()
            query = cgSQL.sql[cgSQL.sqlCode.versionHistory]
            c.execute(query, (self.docType, self.docScope))
        except:
            errorHandler(self.errProc, errorCode.cannotQuery,
                         cgSQL.sqlCode.versionHistory, query, self.docType + ', ' + self.docScope)

        #----------------------------------------------------------------------#
        # Set the document version information into the core properties. The   #
        # last version number is first in the returned descending order cursor:#
        #----------------------------------------------------------------------#
        data = c.fetchone()
        if (data is None):
            pass
        else:
            #------------------------------------------------------------------#
            # Set the document version information into the core properties:   #
            #------------------------------------------------------------------#
            self.docVer = data['Ver']
            self.docDate = data['ChangedDate']
            self.document.core_properties.category = self.docVer
            self.document.core_properties.keywords = self.docDate

    #--------------------------------------------------------------------------#
    # Function: tableAddRows                                                   #
    #                                                                          #
    # Description:                                                             #
    # Adds one table row for each row in the cursor.                           #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string to get the data for.              #
    #--------------------------------------------------------------------------#
    def tableAddRows(self, query):
        #----------------------------------------------------------------------#
        # Define global variables:                                             #
        #----------------------------------------------------------------------#
        global ps

        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableAddRows'

        #----------------------------------------------------------------------#
        # Execute the query and get the data:                                  #
        #----------------------------------------------------------------------#
        bHasData = False
        c = self.getQueryData(query)
        data = c.fetchall()

        #----------------------------------------------------------------------#
        # Iterate the data to get the row count:                               #
        #----------------------------------------------------------------------#
        self.rowCount = 0
        for row in data:
            bHasData = True
            self.rowCount = self.rowCount + 1

        #----------------------------------------------------------------------#
        # Re-query to refresh the cursor:                                      #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        data = c.fetchall()

        #----------------------------------------------------------------------#
        # Get the column attribute row:                                        #
        #----------------------------------------------------------------------#
        iRow = 0
        rowAttr = self.currentTable.rows[1]
        cellsAttr = rowAttr.cells
        para = cellsAttr[0].paragraphs[0]
        styleAttr = para.style

        #----------------------------------------------------------------------#
        # Process each row of returned data:                                   #
        #----------------------------------------------------------------------#
        for row in data:
            #------------------------------------------------------------------#
            # Update the progress bar is not a parent document:
            #------------------------------------------------------------------#
            iRow = iRow + 1
            if (not self.parent):
                #--------------------------------------------------------------#
                # Update the progress bar:                                     #
                #--------------------------------------------------------------#
                pc = 0.8 / self.rowCount
                ps.update(pc)
                ps.set_description('Row ' + str(iRow))
                ps.refresh()

            #------------------------------------------------------------------#
            # Add a new row to the table and enter a loop to process each      #
            # cell:                                                            #
            #------------------------------------------------------------------#
            rowNew = self.currentTable.add_row()
#            rowNew.height_rule = WD_ROW_HEIGHT.EXACTLY
            rowNew.height_rule = 2
            rowNew.height = shared.Cm(1.2)
#            cellsNew = self.currentTable.add_row().cells
            cellsNew = rowNew.cells
            for i in range(0, len(cellsAttr)):
                #--------------------------------------------------------------#
                # Replace the field placeholders in the cell text:             #
                #--------------------------------------------------------------#
                s = cellsAttr[i].text
                for fld in row.keys():
                    s = s.replace('@@' + fld.upper() + '@@', str(row[fld]))
                para = cellsNew[i].paragraphs[0]
                para.text = ''
                run = para.add_run(s)
                para.style = styleAttr

        #----------------------------------------------------------------------#
        # Clean up the table by deleting the query and attribute rows:         #
        #----------------------------------------------------------------------#
        self.remove_row(self.currentTable, rowAttr)
#        self.currentTable.style = 'testTable'

        #----------------------------------------------------------------------#
        # There is no data so write a message:                                 #
        #----------------------------------------------------------------------#
        if (not bHasData):
            cellsNew = self.currentTable.add_row().cells
            a = cellsNew[0]
            for i in range(0, len(cellsNew)):
                b = cellsNew[i]
                a.merge(b)

#            tStyle = self.currentTable.style
            para = cellsNew[0].paragraphs[0]
            para.style = styleAttr
            run = para.add_run('No data to display. Refer to the Functional Specification.')
#            self.currentTable.style = tStyle

    #--------------------------------------------------------------------------#
    # Function: tableBaseRecord                                                #
    #                                                                          #
    # Description:                                                             #
    # A child document table which includes fields from the parent document    #
    # base cursor. The current record data should be inserted.                 #
    #--------------------------------------------------------------------------#
    def tableBaseRecord(self):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of any programming errors:         #
        #----------------------------------------------------------------------#
        self.errProc = 'tableBaseRecord'

        #----------------------------------------------------------------------#
        # Process the table as a static field table with the global data:      #
        #----------------------------------------------------------------------#
        for fld in self.dataRow.keys():
            #------------------------------------------------------------------#
            # Ignore non-ascii characters:                                     #
            #------------------------------------------------------------------#
            try:
                self.dataRow[fld].decode('ascii')
                self.srTable(self.currentTable, '@@' + fld.upper() + '@@', str(self.dataRow[fld]))
            except:
                pass

    #--------------------------------------------------------------------------#
    # Function: tableInsertImage                                               #
    #                                                                          #
    # Description:                                                             #
    # Inserts an image into the table cell.                                    #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # sImage                The image filename.                                #
    #--------------------------------------------------------------------------#
    def tableInsertImage(self, sImage):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableInsertImage'

        #----------------------------------------------------------------------#
        # Get the image file:                                                  #
        #----------------------------------------------------------------------#
        for fld in self.dataRow.keys():
            sImage = sImage.replace('@@' + fld.upper() + '@@', str(self.dataRow[fld]))

        #----------------------------------------------------------------------#
        # Check if the image exists:                                           #
        #----------------------------------------------------------------------#
        sImage = self.inputDir + '/' + sImage
        if not os.path.exists(sImage):
            errorHandler(self.errProc, errorCode.fileNotExist, sImage)

        #----------------------------------------------------------------------#
        # Insert the image:                                                    #
        #----------------------------------------------------------------------#
        para = self.currentRow.cells[0].paragraphs[0]
        style = para.style
        para.text = ''
        run = para.add_run()
        run.add_picture(sImage)
        para.style = style

    #--------------------------------------------------------------------------#
    # Function: tableStaticFields                                              #
    #                                                                          #
    # Description:                                                             #
    # Populates a table with placeholder data without altering the table       #
    # structure.                                                               #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def tableStaticFields(self, query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableStaticFields'

        #----------------------------------------------------------------------#
        # Execute the query and get the data. Static table can only handle     #
        # one record:                                                          #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        data = c.fetchone()
        if (not data is None):
            #------------------------------------------------------------------#
            # Process each field in the data and all cells in the table:       #
            #------------------------------------------------------------------#
            for fld in data.keys():
                #--------------------------------------------------------------#
                # Ignore non-ascii characters:                                 #
                #--------------------------------------------------------------#
                try:
                    data[fld].decode('ascii')
                    self.srTable(self.currentTable, '@@' + fld.upper() + '@@', str(data[fld]))
                except:
                    pass

    #--------------------------------------------------------------------------#
    # Function: getQueryData                                                   #
    #                                                                          #
    # Description:                                                             #
    # Queries the database with the specified SQL string.                      #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryData(self, txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of programming errors:             #
        #----------------------------------------------------------------------#
        self.errProc = 'getQueryData'

        #----------------------------------------------------------------------#
        # Get any parameters from the query string:                            #
        #----------------------------------------------------------------------#
        q = self.getQueryParameters(txtQuery)

        #----------------------------------------------------------------------#
        # Execute the query:                                                   #
        #----------------------------------------------------------------------#
        try:
            c = self.conn.cursor()
            c.execute(q.query, q.parms)
        except:
            errorHandler(self.errProc, errorCode.cannotQuery, q.query, q.parms)
        return c

    #--------------------------------------------------------------------------#
    # Function: getQueryParameters                                             #
    #                                                                          #
    # Description:                                                             #
    # Gets the parameter values from the parameter names in the query string.  #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryParameters(self, txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of any programming errors:         #
        #----------------------------------------------------------------------#
        errProc = 'getQueryParameters'

        #----------------------------------------------------------------------#
        # Declare the named tuple for the returned data:                       #
        #----------------------------------------------------------------------#
        Query = collections.namedtuple('Query', ['query', 'parms'])

        #----------------------------------------------------------------------#
        # Declare local parameter list to return:                              #
        #----------------------------------------------------------------------#
        parms = []

        #----------------------------------------------------------------------#
        # Enter a loop to process all of the parameter placeholders in the     #
        # query:                                                               #
        #----------------------------------------------------------------------#
        while (txtQuery.find('@@') >= 0):
            #------------------------------------------------------------------#
            # Search for any parameter placeholders in the query string:       #
            #------------------------------------------------------------------#
            iTagBegin = txtQuery.find('@@')
            iTagEnd = txtQuery.find('@@', iTagBegin + 2)
            if (iTagBegin == -1):
                #--------------------------------------------------------------#
                # Nothing to do. Return null:                                  #
                #--------------------------------------------------------------#
                return;

            elif (iTagBegin > iTagEnd):
                errorHandler(self.errProc, errorCode.noEndPlaceholder, txtQuery)
            else:
                #--------------------------------------------------------------#
                # Get the parameter text between the BEGIN and END tags and    #
                # replace the parameter with a question mark. Add the          #
                # parameter to the list:                                       #
                #--------------------------------------------------------------#
                sParameterName = txtQuery[iTagBegin + 2:iTagEnd]
                txtQuery = txtQuery[:iTagBegin] + '?' + txtQuery[iTagEnd + 2:]
                parms.append(sParameterName)

        #----------------------------------------------------------------------#
        # Replace the parameters with the existing property values:            #
        #----------------------------------------------------------------------#
        for i in range(len(parms)):
            if (parms[i].upper() == self.fieldKey):
                parms[i] = self.key
            elif (parms[i].upper() == 'SCOPE'):
                parms[i] = self.docScope
            elif (parms[i].upper() == 'DOCTYPE'):
                parms[i] = self.docType

        #----------------------------------------------------------------------#
        # Replace the parameters with the field values in the base if there is #
        # any data:                                                            #
        #----------------------------------------------------------------------#
        c = self.dataBase
        data = None
        try:
            data = c.fetchone()
        except:
            pass
        if (data is not None):
            for i in range(len(parms)):
                for fld in data.keys():
                    if (parms[i].upper() == fld.upper()):
                        parms[i] = data[fld]

        #----------------------------------------------------------------------#
        # Return the parameter value list:                                     #
        #----------------------------------------------------------------------#
        q = Query(query=txtQuery, parms=parms)
        return q;

    def srParagraph(self, paragraph, txtSearch, txtReplace):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'srParagraph'

        #----------------------------------------------------------------------#
        # Update the text in the paragraph by replacing the search text:       #
        #----------------------------------------------------------------------#
        s = paragraph.text
        if (s.find(txtSearch) >= 0):
            style = paragraph.style
            s = s.replace(txtSearch, txtReplace)
            paragraph.text = ''
            run = paragraph.add_run(s)
            paragraph.style = style

    def srDocument(self, document, txtSearch, txtReplace):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'srDocument'

        for paragraph in self.document.paragraphs:
            self.srParagraph(paragraph, txtSearch, txtReplace)

    def srHeader(self, document, txtSearch, txtReplace):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'srHeader'

        for section in self.document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self.srParagraph(paragraph, txtSearch, txtReplace)

    def srTable(self, table, txtSearch, txtReplace):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'srTable'

        for row in self.currentTable.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.srParagraph(paragraph, txtSearch, txtReplace)

    def remove_row(self, t, r):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'remove_row'

        tbl = t._tbl
        tr = r._tr
        tbl.remove(tr)

    #--------------------------------------------------------------------------#
    # Function: printPDF                                                       #
    #                                                                          #
    # Description:                                                             #
    # Prints a docx document to a PDF document using libreoffice writer.       #
    #--------------------------------------------------------------------------#
    def printPDF(self):
        #----------------------------------------------------------------------#
        # Save the output document as PDF:                                     #
        #----------------------------------------------------------------------#
        output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf' ,
                                         '--outdir', self.outputDir, self.outputFile])
#        f = open(os.devnull, 'w')
#        sys.stdout = f
        with silence_stdout():
            print(output)
        self.outputPDFFileName = self.outputDir + '/' + self.outputBaseName + '.pdf'

    #--------------------------------------------------------------------------#
    # Function: docxDelete                                                     #
    #                                                                          #
    # Description:                                                             #
    # Deletes the output docx document. Not needed if want only PDF.           #
    #--------------------------------------------------------------------------#
    def docxDelete(self):
        os.remove(self.outputFile)

    #--------------------------------------------------------------------------#
    # Function: tagRenumber                                                    #
    #                                                                          #
    # Description:                                                             #
    # Renumbers the tag list.                                                  #
    #--------------------------------------------------------------------------#
    def tagRenumber(self):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'tagRenumber'

        #----------------------------------------------------------------------#
        # Use global tag number variable:                                      #
        #----------------------------------------------------------------------#
        global iRefNum
        global sRefPrefix

        #----------------------------------------------------------------------#
        # Loop through all the tables in the document table collection:        #
        #----------------------------------------------------------------------#
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #------------------------------------------------------#
                        # Check if a valid tag number:                         #
                        #------------------------------------------------------#
                        s = paragraph.text
                        if (len(s) >= 1 + len(sRefPrefix) and s[:len(sRefPrefix)] == sRefPrefix):
#                            pa = re.compile('[a-zA-Z]')
                            pn = re.compile('[0-9]')
#                            ma = pa.match(s[:len(sRefPrefix)])
                            mn = pn.match(s[len(sRefPrefix):])
                            if (not mn is None):
                                style = paragraph.style
                                paragraph.text = ''
                                sn = sRefPrefix + str(iRefNum)
                                run = paragraph.add_run(sn)
                                paragraph.style = style
                                iRefNum = iRefNum + 1

#------------------------------------------------------------------------------#
# Class: gDocParent                                                            #
#                                                                              #
# Description:                                                                 #
# Parent document object which may have child documents appended.              #
#------------------------------------------------------------------------------#
class gDocParent(gDoc):
    #--------------------------------------------------------------------------#
    # Function: getParentData                                                  #
    #                                                                          #
    # Description:                                                             #
    # Processes all of the tables in the document looking for the base data    #
    # definition query. Need to populate first before any other table.         #
    #--------------------------------------------------------------------------#
    def getParentData(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'getParentData'

        #----------------------------------------------------------------------#
        # Initialise the tag number at 1:                                      #
        #----------------------------------------------------------------------#
        iRefNum = 1

        #----------------------------------------------------------------------#
        # Process each table in the document:                                  #
        #----------------------------------------------------------------------#
        for table in self.document.tables:
            #------------------------------------------------------------------#
            # Get the table data source:                                       #
            #------------------------------------------------------------------#
            txtQuery = self.currentRow.cells[0].text
            txtQuery = txtQuery.upper()

            #------------------------------------------------------------------#
            # Check for non-ascii characters. Can't be a keyword string in     #
            # that case:                                                       #
            #------------------------------------------------------------------#
            try:
                txtQuery.decode('ascii')
            except:
                #--------------------------------------------------------------#
                # Just ignore this table:                                      #
                #--------------------------------------------------------------#
                pass
            else:
                #--------------------------------------------------------------#
                # Check if base query data for the entire document:            #
                #--------------------------------------------------------------#
                if (txtQuery[:7] == 'SQLBASE'):
                    #----------------------------------------------------------#
                    # Get the base data for the document:                      #
                    #----------------------------------------------------------#
                    self.getBaseData(txtQuery[8:])
                else:
                    #----------------------------------------------------------#
                    # Ignore other content:                                    #
                    #----------------------------------------------------------#
                    pass

    #--------------------------------------------------------------------------#
    # Function: appendPDF                                                      #
    #                                                                          #
    # Description:                                                             #
    # Appends one PDF file to the end of another.                              #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # dc                    The child document object.                         #
    #--------------------------------------------------------------------------#
    def appendPDF(self, dc):
        #----------------------------------------------------------------------#
        # Open the parent and child PDF documents:                             #
        #----------------------------------------------------------------------#
        pdfParent = open(self.outputPDFFileName, 'rb')
        pdfChild = open(dc.outputPDFFileName, 'rb')

        #----------------------------------------------------------------------#
        # Create a new PDF file merger object:                                 #
        #----------------------------------------------------------------------#
        merger = PyPDF2.PdfFileMerger()

        #----------------------------------------------------------------------#
        # Append the child PDF content to the parent PDF document and save in  #
        # a new temporary document in the output directory:                    #
        #----------------------------------------------------------------------#
        merger.append(fileobj=pdfParent)
        merger.append(fileobj=pdfChild)
        sOutputPDF = self.outputDir + '/p.pdf'
        merger.write(open(sOutputPDF, 'wb'))

        #----------------------------------------------------------------------#
        # Put the output file back to the main PDF output and delete the child #
        # record and other tepmorary files:                                    #
        #----------------------------------------------------------------------#
        os.remove(self.outputPDFFileName)
        copyfile(sOutputPDF, self.outputPDFFileName)
        os.remove(sOutputPDF)
        os.remove(dc.outputPDFFileName)

#------------------------------------------------------------------------------#
# Class: gDocChild                                                             #
#                                                                              #
# Description:                                                                 #
# Child document object which must be appended to its parent.                  #
#------------------------------------------------------------------------------#
class gDocChild(gDoc):
    #--------------------------------------------------------------------------#
    # Constructor:                                                             #
    #--------------------------------------------------------------------------#
    def __init__(self, p, sInputChild, data, sFieldKey, sFieldKeyValue, iChildNum):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'childInit'

        #----------------------------------------------------------------------#
        # Set the instance attributes:                                         #
        #----------------------------------------------------------------------#
        self.inputFile = sInputChild
        self.childNum = iChildNum
        self.dataBase = p.dataBase
        self.dataRow = data
        self.key = sFieldKeyValue
        self.fieldKey = sFieldKey.upper()
        self.conn = p.conn
        self.docType = p.docType
        self.docScope = p.docScope

        #----------------------------------------------------------------------#
        # Set the file characteristics:                                        #
        #----------------------------------------------------------------------#
        self.outputDir = p.outputDir
        self.outputFile = self.outputDir + '/t.docx'
        self.outputFileName = os.path.basename(self.outputFile)
        self.outputBaseName = os.path.splitext(self.outputFileName)[0]

        #----------------------------------------------------------------------#
        # Make sure the input document file exists:                            #
        #----------------------------------------------------------------------#
        self.inputDir = os.path.dirname(self.inputFile)
        if not os.path.exists(self.inputFile):
            errorHandler(self.errProc, errorCode.fileChildNotExist, self.inputFile)

        #----------------------------------------------------------------------#
        # Create the document:                                                 #
        #----------------------------------------------------------------------#
#        sKey = str(self.dataRow[self.fieldKey])
#        print(sKey)
        self.createDocument()

@contextmanager
def silence_stdout():
    new_target = open(os.devnull, "w")
    old_target, sys.stdout = sys.stdout, new_target
    try:
        yield new_target
    finally:
        sys.stdout = old_target

def pdfPageNumbers(fileName):
    global ps

    f = 't.pdf'
    os.rename(fileName, f)
    pdfFile = open(f, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFile)
    pdfWriter = PyPDF2.PdfFileWriter()

    tmp = '_tmp.pdf'
    n = pdfReader.numPages

    for i in range(n):
        c = canvas.Canvas(tmp)
        c.setFont('Helvetica', 7)
#        c.drawString((210//2)*mm, (4)*mm, 'Page ' + str(i + 1) + ' of ' + str(n))
        c.drawString((180)*mm, (276)*mm, 'Page ' + str(i + 1) + ' of ' + str(n))
        c.drawString((273)*mm, (197.5)*mm, 'Page ' + str(i + 1) + ' of ' + str(n))
        c.showPage()
        c.save()
        pageObj = pdfReader.getPage(i)
#        cObj = canvas.Canvas(pageObj)
#        print(cObj._pagesize)
#        page_width, page_height = cObj._pagesize
#        print('width ' + str(page_width))
#        print('height ' + str(page_height))
        pdfTmp = open(tmp, 'rb')
        pdfTmpReader = PyPDF2.PdfFileReader(pdfTmp)
        pageObj.mergePage(pdfTmpReader.getPage(0))
        os.remove(tmp)
        pdfWriter.addPage(pageObj)

        #----------------------------------------------------------------------#
        # Update the progress bar:                                             #
        #----------------------------------------------------------------------#
        pc = 0.2 / n
        ps.update(pc)
        ps.set_description('Numbering page ' + str(i))
        ps.refresh()

    pdfOutputFile = open(fileName, 'wb')
    pdfWriter.write(pdfOutputFile)
    pdfOutputFile.close()
    pdfFile.close()
    os.remove(f)

#------------------------------------------------------------------------------#
# Call the main function:                                                      #
#------------------------------------------------------------------------------#
main()
