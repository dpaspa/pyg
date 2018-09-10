#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Rev By               Date        CC        Note                              #
# 1.0 David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import argparse
from enum import Enum
#from docx.document import Document
from docx import Document
from docx import shared
from docx.oxml.shared import OxmlElement, qn # Necessary Import
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os.path
from shutil import copyfile
import sys
import traceback
from tqdm import trange
from time import sleep
import sqlite3
import collections
import datetime
from contextlib import contextmanager
import cgSQL

from dgProperty import setProperty

import logging
logging.basicConfig(level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Declare global variables:                                                    #
#------------------------------------------------------------------------------#
currentRow = ''
currentTable = ''

#------------------------------------------------------------------------------#
# Declare the application title and calling arguments help:                    #
#------------------------------------------------------------------------------#
appTitle = 'Document Generator'
appVersion = '1'
parser = argparse.ArgumentParser(description='Generates a document from a configuration spreadsheet and set of document templates')
parser.add_argument('-c','--config', help='Configuration database', required=True)
parser.add_argument('-i','--input', help='Input base document template file', required=True)
parser.add_argument('-o','--output', help='Output for the generated document file', required=True)
parser.add_argument('-f','--filter', help='The database WHERE clause filter expression', required=False)
args = vars(parser.parse_args())

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Close the progress bar:                                                  #
    #--------------------------------------------------------------------------#
    ps.close()

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print(appTitle + ' Version ' + appVersion + '\r\n' + 'ERROR ' +
          str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
#    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    cannotConnectDB                    = -1
    cannotGetSQL                       = -2
    cannotQuery                        = -3
    fileNotExist                       = -4
    pathNotExist                       = -5
    noEndPlaceholder                   = -6
    nonASCIICharacter                  = -7
    nonASCIIField                      = -8
    unknownAttribute                   = -9

errorMessage = {
    errorCode.cannotConnectDB          : 'Cannot connect to sqlite database @1',
    errorCode.cannotGetSQL             : 'Cannot retrieve SQL query expression for attribute @1',
    errorCode.cannotQuery              : 'Query cannot execute using SQL expression @1 with parameters @2',
    errorCode.fileNotExist             : 'Parent file @1 does not exist.',
    errorCode.pathNotExist             : 'Path @1 does not exist.',
    errorCode.noEndPlaceholder         : 'No valid END placholder in query string @1.',
    errorCode.nonASCIICharacter        : 'Query expression @1 returned a non ascii-encoded unicode string',
    errorCode.nonASCIIField            : 'Query expression @1 referenced a non ascii-encoded field @2',
    errorCode.unknownAttribute         : 'Attribute @1 is unknown.'
}

#------------------------------------------------------------------------------#
# Create a progress bar:                                                       #
#------------------------------------------------------------------------------#
ps = trange(1, desc='Create document...', leave=False)

#------------------------------------------------------------------------------#
# Function main                                                                #
#                                                                              #
# Description:                                                                 #
# Creates a generated document using sqlite data to populate a document        #
# template which has @@ placeholders in it. The document must be listed in the #
# database in a tblDocument list.                                              #
#                                                                              #
# THe template is filled with data and saved as a new docx file and is also    #
# output as a PDF file.                                                        #
#                                                                              #
# The document template must have the query definitions defined within it in   #
# tables with the query string in the first row. Different types of queries    #
# are handled as follows:                                                      #
#                                                                              #
# IMAGE                 An image will be inserted in the table cell.           #
# SQLBASE               Base data for the document, such as document number    #
#                       and title.                                             #
# SQLRECORD             A sub-document in a multi-record compound document.    #
#                       The first document will be the parent header document  #
#                       and each record will be appended as a child document.  #
# SQLROW                Multi-record table. The table must include a row of    #
#                       column headings followed by a row of cell data         #
#                       attributes using @@ data field placeholders.           #
# SQLSTATIC             A table of any cell arrangement to be populated        #
#                       without changing the table structure.                  #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# d                     The generic document object.                           #
#------------------------------------------------------------------------------#
def dataDocument(d):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = dataDocument.__name__

    #--------------------------------------------------------------------------#
    # Use global variables:                                                    #
    #--------------------------------------------------------------------------#
    global ps

    #--------------------------------------------------------------------------#
    # Enter a loop to process each table in the document:                      #
    #--------------------------------------------------------------------------#
    numTables = len(d.document.tables)
    iTable = 0
    for table in d.document.tables:
        #----------------------------------------------------------------------#
        # Process the table and update any SQL data:                           #
        #----------------------------------------------------------------------#
        iTable = iTable + 1
        dataTable(table)

        #----------------------------------------------------------------------#
        # Update the progress bar:                                             #
        #----------------------------------------------------------------------#
        pc = 1.0 / numTables
        ps.update(pc)
        ps.set_description('Updating data in table ' + str(iTable))
        ps.refresh()

    #--------------------------------------------------------------------------#
    # Report completion regardless of error:                                   #
    #--------------------------------------------------------------------------#
    ps.set_description('Processing complete')
    ps.refresh()
    ps.close()

    #--------------------------------------------------------------------------#
    # Close the database and output a success message:                         #
    #--------------------------------------------------------------------------#
    conn.close()
    print('Congratulations! Document ' + sOutput + ' generated successfully.')

    #--------------------------------------------------------------------------#
    # Function: dataTable                                                      #
    #                                                                          #
    # Description:                                                             #
    # Processes the current table looking for defined header keywords defining #
    # the content of the table.                                                #
    #--------------------------------------------------------------------------#
    def dataTable(table):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        errProc = 'dataTable'

        #----------------------------------------------------------------------#
        # Use global variables:                                                #
        #----------------------------------------------------------------------#
        global currentRow
        global currentTables

        #----------------------------------------------------------------------#
        # Get the table data source:                                           #
        #----------------------------------------------------------------------#
        bDelRow = False
        currentTable = table
        currentRow = currentTable.rows[0]
        txtQuery = currentRow.cells[0].text

        #----------------------------------------------------------------------#
        # Check for non-ascii characters. Can't be a keyword string in that    #
        # case:                                                                #
        #----------------------------------------------------------------------#
        try:
            txtQuery.decode('ascii')
        except:
            errorHandler(errProc, errorCode.nonASCIICharacter, txtQuery)
        else:
            #------------------------------------------------------------------#
            # Check for the content type. Check if an image anchor:            #
            #------------------------------------------------------------------#
            if (txtQuery[:5] == 'IMAGE'):
                #--------------------------------------------------------------#
                # Insert the image:                                            #
                #--------------------------------------------------------------#
                remove_row(currentTable, currentRow)
                currentRow = currentTable.rows[0]
                tableInsertImage(txtQuery[6:])

            #------------------------------------------------------------------#
            # Check if a documentn property:                                   #
            #------------------------------------------------------------------#
            elif (txtQuery[:7] == 'SQLPROP'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                remove_row(currentTable, currentRow)
                docProperty(txtQuery[8:])

            #------------------------------------------------------------------#
            # Check if an append table rows query:                             #
            #------------------------------------------------------------------#
            elif (txtQuery[:6] == 'SQLROW'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                remove_row(currentTable, currentRow)
                currentRow = currentTable.rows[0]
                tableAddRows(txtQuery[7:])

            #------------------------------------------------------------------#
            # Check if a static table query:                                   #
            #------------------------------------------------------------------#
            elif (txtQuery[:9] == 'SQLSTATIC'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                remove_row(currentTable, currentRow)
                currentRow = currentTable.rows[0]
                tableStaticFields(txtQuery[10:])

            #------------------------------------------------------------------#
            # Check if an append table rows query:                             #
            #------------------------------------------------------------------#
            elif (txtQuery[:10] == 'SQLVERHIST'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                remove_row(currentTable, currentRow)
                currentRow = currentTable.rows[0]
                query = cgSQL.sql[cgSQL.sqlCode.VERHIST]
                tableAddRows(query)
            else:
                #--------------------------------------------------------------#
                # Ignore other content:                                        #
                #--------------------------------------------------------------#
                pass

    #--------------------------------------------------------------------------#
    # Function: docProperty                                                    #
    #                                                                          #
    # Description:                                                             #
    # Sets the document property based on the query result.                    #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def docProperty( query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        errProc = 'docProperty'

        #----------------------------------------------------------------------#
        # Execute the query and get the data. Static table can only handle     #
        # one record:                                                          #
        #----------------------------------------------------------------------#
        c = getQueryData(query)
        data = c.fetchone()
        if (not data is None):
            #------------------------------------------------------------------#
            # Try to set the document property if the field exists:            #
            #------------------------------------------------------------------#
            try:
                prop = data.keys()
                setProperty(document, prop[0], data[0])
            except:
                pass

    #--------------------------------------------------------------------------#
    # Function: tableAddRows                                                   #
    #                                                                          #
    # Description:                                                             #
    # Adds one table row for each row in the cursor.                           #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string to get the data for.              #
    #--------------------------------------------------------------------------#
    def tableAddRows( query):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'tableAddRows'

        #----------------------------------------------------------------------#
        # Execute the query and get the data:                                  #
        #----------------------------------------------------------------------#
        bHasData = False
        c = getQueryData(query)
        data = c.fetchall()

        #----------------------------------------------------------------------#
        # Iterate the data to get the row count:                               #
        #----------------------------------------------------------------------#
        rowCount = 0
        for row in data:
            bHasData = True
            rowCount = rowCount + 1

        #----------------------------------------------------------------------#
        # There is no data so delete the entire table and the paragraph:       #
        #----------------------------------------------------------------------#
        if (not bHasData):
            p1 = getTableParagraph(currentTable)
            remove_table(currentTable)
            remove_paragraph(p1)
        else:
            #------------------------------------------------------------------#
            # Re-query to refresh the cursor:                                  #
            #------------------------------------------------------------------#
            c = getQueryData(query)
            data = c.fetchall()

            #------------------------------------------------------------------#
            # Get the column attribute row:                                    #
            #------------------------------------------------------------------#
            rowAttr = currentTable.rows[1]
            cellsAttr = rowAttr.cells
            para = cellsAttr[0].paragraphs[0]
            styleAttr = para.style

            #------------------------------------------------------------------#
            # Process each row of returned data:                               #
            #------------------------------------------------------------------#
            for row in data:
                #--------------------------------------------------------------#
                # Add a new row to the table and enter a loop to process each  #
                # cell:                                                        #
                #--------------------------------------------------------------#
                rowNew = currentTable.add_row()
    #            rowNew.height_rule = WD_ROW_HEIGHT.EXACTLY
                rowNew.height_rule = 2
                rowNew.height = shared.Cm(1.2)
    #            cellsNew = currentTable.add_row().cells
                cellsNew = rowNew.cells
                for i in range(0, len(cellsAttr)):
                    #----------------------------------------------------------#
                    # Replace the field placeholders in the cell text:         #
                    #----------------------------------------------------------#
                    s = cellsAttr[i].text
                    for fld in row.keys():
                        try:
                            s = s.replace('@@' + fld.upper() + '@@', str(row[fld]))
                        except:
                            pass
                    para = cellsNew[i].paragraphs[0]
                    para.text = ''
                    run = para.add_run(s)
                    para.style = styleAttr

            #------------------------------------------------------------------#
            # Clean up the table by deleting the query and attribute rows:     #
            #------------------------------------------------------------------#
            remove_row(currentTable, rowAttr)

    #--------------------------------------------------------------------------#
    # Function: tableInsertImage                                               #
    #                                                                          #
    # Description:                                                             #
    # Inserts an image into the table cell.                                    #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # sImage                The image filename.                                #
    #--------------------------------------------------------------------------#
    def tableInsertImage( sImage):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        errProc = 'tableInsertImage'

        #----------------------------------------------------------------------#
        # Get the image file:                                                  #
        #----------------------------------------------------------------------#
        for fld in dataRow.keys():
            sImage = sImage.replace('@@' + fld.upper() + '@@', str(dataRow[fld]))

        #----------------------------------------------------------------------#
        # Check if the image exists:                                           #
        #----------------------------------------------------------------------#
        sImage = inputDir + '/' + sImage
        if not os.path.exists(sImage):
            errorHandler(errProc, errorCode.fileNotExist, sImage)

        #----------------------------------------------------------------------#
        # Insert the image:                                                    #
        #----------------------------------------------------------------------#
        para = currentRow.cells[0].paragraphs[0]
        style = para.style
        para.text = ''
        run = para.add_run()
        run.add_picture(sImage)
        para.style = style

    #--------------------------------------------------------------------------#
    # Function: tableStaticFields                                              #
    #                                                                          #
    # Description:                                                             #
    # Populates a table with placeholder data without altering the table       #
    # structure.                                                               #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def tableStaticFields( query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        errProc = 'tableStaticFields'

        #----------------------------------------------------------------------#
        # Execute the query and get the data. Static table can only handle     #
        # one record:                                                          #
        #----------------------------------------------------------------------#
        c = getQueryData(query)
        data = c.fetchone()
        if (data is None):
            #------------------------------------------------------------------#
            # No data so delete the table and the paragraph:                   #
            #------------------------------------------------------------------#
            p1 = getTableParagraph(currentTable)
            remove_table(currentTable)
            remove_paragraph(p1)
        else:
            #------------------------------------------------------------------#
            # Process each field in the data and all cells in the table:       #
            #------------------------------------------------------------------#
            for fld in data.keys():
                #--------------------------------------------------------------#
                # Ignore non-ascii characters:                                 #
                #--------------------------------------------------------------#
                try:
                    str(data[fld]).decode('ascii')
                    srTable(currentTable, '@@' + fld.upper() + '@@', str(data[fld]))
                except:
                    errorHandler(errProc, errorCode.nonASCIIField, query, fld)

    #--------------------------------------------------------------------------#
    # Function: getQueryData                                                   #
    #                                                                          #
    # Description:                                                             #
    # Queries the database with the specified SQL string.                      #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryData( txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of programming errors:             #
        #----------------------------------------------------------------------#
        errProc = 'getQueryData'

        #----------------------------------------------------------------------#
        # Get any parameters from the query string:                            #
        #----------------------------------------------------------------------#
        q = getQueryParameters(txtQuery)

        #----------------------------------------------------------------------#
        # Execute the query:                                                   #
        #----------------------------------------------------------------------#
        try:
            c = conn.cursor()
            c.execute(q.query, q.parms)
        except:
            errorHandler(errProc, errorCode.cannotQuery, q.query, q.parms)
        return c

    #--------------------------------------------------------------------------#
    # Function: getQueryParameters                                             #
    #                                                                          #
    # Description:                                                             #
    # Gets the parameter values from the parameter names in the query string.  #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryParameters( txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of any programming errors:         #
        #----------------------------------------------------------------------#
        errProc = 'getQueryParameters'

        #----------------------------------------------------------------------#
        # Declare the named tuple for the returned data:                       #
        #----------------------------------------------------------------------#
        Query = collections.namedtuple('Query', ['query', 'parms'])

        #----------------------------------------------------------------------#
        # Declare local parameter list to return:                              #
        #----------------------------------------------------------------------#
        parms = []

        #----------------------------------------------------------------------#
        # Enter a loop to process all of the parameter placeholders in the     #
        # query:                                                               #
        #----------------------------------------------------------------------#
        while (txtQuery.find('@@') >= 0):
            #------------------------------------------------------------------#
            # Search for any parameter placeholders in the query string:       #
            #------------------------------------------------------------------#
            iTagBegin = txtQuery.find('@@')
            iTagEnd = txtQuery.find('@@', iTagBegin + 2)
            if (iTagBegin == -1):
                #--------------------------------------------------------------#
                # Nothing to do. Return null:                                  #
                #--------------------------------------------------------------#
                return;

            elif (iTagBegin > iTagEnd):
                errorHandler(errProc, errorCode.noEndPlaceholder, txtQuery)
            else:
                #--------------------------------------------------------------#
                # Get the parameter text between the BEGIN and END tags and    #
                # replace the parameter with a question mark. Add the          #
                # parameter to the list:                                       #
                #--------------------------------------------------------------#
                sParameterName = txtQuery[iTagBegin + 2:iTagEnd]
                txtQuery = txtQuery[:iTagBegin] + '?' + txtQuery[iTagEnd + 2:]
                parms.append(sParameterName)

        #----------------------------------------------------------------------#
        # Replace the parameters with the existing property values:            #
        #----------------------------------------------------------------------#
        for i in range(len(parms)):
            if (parms[i].upper() == 'FILTER'):
                parms[i] = filter

        #----------------------------------------------------------------------#
        # Return the parameter value list:                                     #
        #----------------------------------------------------------------------#
        q = Query(query=txtQuery, parms=parms)
        return q;

    def srParagraph( paragraph, txtSearch, txtReplace):
        errProc = 'srParagraph'

        #----------------------------------------------------------------------#
        # Update the text in the paragraph by replacing the search text:       #
        #----------------------------------------------------------------------#
        s = paragraph.text
        if (s.find(txtSearch) >= 0):
            style = paragraph.style
#            txtReplace = txtReplace.replace('\"','')
            s = s.replace(txtSearch, txtReplace)
            paragraph.text = ''
            run = paragraph.add_run(s)
            paragraph.style = style

    def srDocument( document, txtSearch, txtReplace):
        errProc = 'srDocument'

        for paragraph in document.paragraphs:
            srParagraph(paragraph, txtSearch, txtReplace)

    def srHeader( document, txtSearch, txtReplace):
        errProc = 'srHeader'

        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                srParagraph(paragraph, txtSearch, txtReplace)

    def srTable( table, txtSearch, txtReplace):
        errProc = 'srTable'

        for row in currentTable.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    srParagraph(paragraph, txtSearch, txtReplace)

    def remove_paragraph( paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def remove_row( t, r):
        errProc = 'remove_row'
        tbl = t._tbl
        tr = r._tr
        tbl.remove(tr)

    def remove_table( t):
        errProc = 'remove_table'
        tbl = currentTable._tbl
        for row in currentTable.rows:
            tr = row._tr
            tbl.remove(tr)
#        tbl = t._tbl
#        document.tables.remove(tbl)

    #--------------------------------------------------------------------------#
    # Return a newly created paragraph, inserted directly before this          #
    # item (Table, etc.):                                                      #
    #--------------------------------------------------------------------------#
#    def insert_paragraph_before( item, text, style=None):
#        p = CT_P.add_p_before(item._element)
#        p2 = Paragraph(p, item._parent)
#        p2.text = text
#        p2.style = style
#        return p2

    #--------------------------------------------------------------------------#
    # Yield each paragraph and table child within *parent*, in document order. #
    # Each returned value is an instance of either Table or Paragraph. *parent*#
    # would most commonly be a reference to a main Document object, but        #
    # also works for a _Cell object, which itcan contain paragraphs and   #
    # tables:                                                                  #
    #--------------------------------------------------------------------------#
    def iter_block_items( parent):
        if isinstance(parent, type(document)):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def getTableParagraph( t):
        nodePrev = document.paragraphs[0]

        for node in iter_block_items(document):
            if isinstance(node, Table):
                if (node._tbl == t._tbl):
                    return nodePrev
            else:
                nodePrev = node
