#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Rev By               Date        CC        Note                              #
# 1.0 David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import argparse
from enum import Enum
#from docx.document import Document
from docx import Document
from docx import shared
from docx.oxml.shared import OxmlElement, qn # Necessary Import
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os.path
from shutil import copyfile
import sys
import traceback
from tqdm import trange
from time import sleep
import sqlite3
import collections
import datetime
from contextlib import contextmanager
import cgSQL

import logging
logging.basicConfig(level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Declare the application title and calling arguments help:                    #
#------------------------------------------------------------------------------#
appTitle = 'Document Generator'
appVersion = '1'
parser = argparse.ArgumentParser(description='Generates a document from a configuration spreadsheet and set of document templates')
parser.add_argument('-c','--config', help='Configuration database', required=True)
parser.add_argument('-i','--input', help='Input base document template file', required=True)
parser.add_argument('-o','--output', help='Output for the generated document file', required=True)
parser.add_argument('-f','--filter', help='The database WHERE clause filter expression', required=False)
args = vars(parser.parse_args())

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Close the progress bar:                                                  #
    #--------------------------------------------------------------------------#
    ps.close()

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print(appTitle + ' Version ' + appVersion + '\r\n' + 'ERROR ' +
          str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
#    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    cannotConnectDB                    = -1
    cannotGetSQL                       = -2
    cannotQuery                        = -3
    fileNotExist                       = -4
    pathNotExist                       = -5
    noEndPlaceholder                   = -6
    nonASCIICharacter                  = -7
    nonASCIIField                      = -8
    unknownAttribute                   = -9

errorMessage = {
    errorCode.cannotConnectDB          : 'Cannot connect to sqlite database @1',
    errorCode.cannotGetSQL             : 'Cannot retrieve SQL query expression for attribute @1',
    errorCode.cannotQuery              : 'Query cannot execute using SQL expression @1 with parameters @2',
    errorCode.fileNotExist             : 'Parent file @1 does not exist.',
    errorCode.pathNotExist             : 'Path @1 does not exist.',
    errorCode.noEndPlaceholder         : 'No valid END placholder in query string @1.',
    errorCode.nonASCIICharacter        : 'Query expression @1 returned a non ascii-encoded unicode string',
    errorCode.nonASCIIField            : 'Query expression @1 referenced a non ascii-encoded field @2',
    errorCode.unknownAttribute         : 'Attribute @1 is unknown.'
}

#------------------------------------------------------------------------------#
# Create a progress bar:                                                       #
#------------------------------------------------------------------------------#
ps = trange(1, desc='Create document...', leave=False)

#------------------------------------------------------------------------------#
# Function main                                                                #
#                                                                              #
# Description:                                                                 #
# The main entry point for the program.                                        #
#------------------------------------------------------------------------------#
def main():
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = main.__name__

    #--------------------------------------------------------------------------#
    # Use the global prefix:                                                   #
    #--------------------------------------------------------------------------#
    global ps
    global sRefPrefix

    #--------------------------------------------------------------------------#
    # Get the input file and output generated document file path and name:     #
    #--------------------------------------------------------------------------#
    sInput = args['input']
    sOutput = args['output']

    #--------------------------------------------------------------------------#
    # Get the database query filter expression:                                #
    #--------------------------------------------------------------------------#
    sFilter = args['filter']

    #--------------------------------------------------------------------------#
    # Get the configuration data database name and check it exists:            #
    #--------------------------------------------------------------------------#
    dbName = args['config']
    if not os.path.exists(dbName):
        errorHandler(errProc, errorCode.filenotExist, dbName)

    #--------------------------------------------------------------------------#
    # Try to connect to the sqlite database file:                              #
    #--------------------------------------------------------------------------#
    conn = ''
    try:
        conn = sqlite3.connect(dbName)
        conn.row_factory = sqlite3.Row
    except:
        errorHandler(errProc, errorCode.cannotConnectDB, dbName)

    #--------------------------------------------------------------------------#
    # Create the new document and populate with data:                          #
    #--------------------------------------------------------------------------#
    d = gDoc(conn, sInput, sOutput, sFilter)

    #--------------------------------------------------------------------------#
    # Report completion regardless of error:                                   #
    #--------------------------------------------------------------------------#
    ps.set_description('Processing complete')
    ps.refresh()
    ps.close()

    #--------------------------------------------------------------------------#
    # Close the database and output a success message:                         #
    #--------------------------------------------------------------------------#
    conn.close()
    print('Congratulations! Document ' + sOutput + ' generated successfully.')

#------------------------------------------------------------------------------#
# Class: gDoc                                                                  #
#                                                                              #
# Description:                                                                 #
# Creates a generated document using sqlite data to populate a document        #
# template which has @@ placeholders in it. The document must be listed in the #
# database in a tblDocument list.                                              #
#                                                                              #
# THe template is filled with data and saved as a new docx file and is also    #
# output as a PDF file.                                                        #
#                                                                              #
# The document template must have the query definitions defined within it in   #
# tables with the query string in the first row. Different types of queries    #
# are handled as follows:                                                      #
#                                                                              #
# IMAGE                 An image will be inserted in the table cell.           #
# SQLBASE               Base data for the document, such as document number    #
#                       and title.                                             #
# SQLRECORD             A sub-document in a multi-record compound document.    #
#                       The first document will be the parent header document  #
#                       and each record will be appended as a child document.  #
# SQLROW                Multi-record table. The table must include a row of    #
#                       column headings followed by a row of cell data         #
#                       attributes using @@ data field placeholders.           #
# SQLSTATIC             A table of any cell arrangement to be populated        #
#                       without changing the table structure.                  #
#------------------------------------------------------------------------------#
# Calling Attributes:                                                          #
# conn                  The database connection object.                        #
# inputFile             The input template file name to use for the document.  #
# outputFile            The docx file name to create after field replacement.  #
# filter                The data filter for the document.                      #
#                                                                              #
# Other Attributes created by the constructor:                                 #
# outputDir             The output directory.                                  #
# outputBaseName        The base output file name without the path.            #
# outputFileName        The base output file name without the extension.       #
#------------------------------------------------------------------------------#
class gDoc(object):
    #--------------------------------------------------------------------------#
    # Constructor:                                                             #
    #--------------------------------------------------------------------------#
    def __init__(self, conn, inputFile, outputFile, filter):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'parentInit'

        #----------------------------------------------------------------------#
        # Make sure the input document file exists:                            #
        #----------------------------------------------------------------------#
        if not os.path.exists(inputFile):
            errorHandler(self.errProc, errorCode.fileNotExist, inputFile)

        #----------------------------------------------------------------------#
        # Set the instance attributes:                                         #
        #----------------------------------------------------------------------#
        self.conn = conn
        self.filter = filter
        self.inputFile = inputFile
        self.outputFile = outputFile

        #----------------------------------------------------------------------#
        # Set the file characteristics:                                        #
        #----------------------------------------------------------------------#
        self.outputFileName = os.path.basename(outputFile)
        self.outputBaseName = os.path.splitext(self.outputFileName)[0]
        self.outputDir = os.path.dirname(self.outputFile)
        if not os.path.exists(self.outputDir):
            errorHandler(self.errProc, errorCode.pathNotExist, self.outputDir)

        #----------------------------------------------------------------------#
        # Create the document:                                                 #
        #----------------------------------------------------------------------#
        self.createDocument()

    #--------------------------------------------------------------------------#
    # Function: createDocument                                                 #
    #                                                                          #
    # Description:                                                             #
    # Creates the docx document and sets its properties and then loops through #
    # all the tables to populate them with data. Finally it saves the          #
    # completed document as both a .docx file and a .pdf file.                 #
    #--------------------------------------------------------------------------#
    def createDocument(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'createDocument'

        #----------------------------------------------------------------------#
        # Open the input document:                                             #
        #----------------------------------------------------------------------#
        self.document = Document(self.inputFile)

        #----------------------------------------------------------------------#
        # Proess all of the tables in the document:                            #
        #----------------------------------------------------------------------#
        self.processTables()

        #----------------------------------------------------------------------#
        # Save the output document:                                            #
        #----------------------------------------------------------------------#
        self.document.save(self.outputFile)

    #--------------------------------------------------------------------------#
    # Function: setProperty                                                    #
    #                                                                          #
    # Description:                                                             #
    # Sets the core document property.                                         #
    #--------------------------------------------------------------------------#
    def comments(self, val):
        self.document.core_properties.comments = val

    def keywords(self, val):
        self.document.core_properties.keywords = val

    def subject(self, val):
        self.document.core_properties.subject = val

    def title(self, val):
        self.document.core_properties.title = val

    def setProperty(self, argument, val):
        switcher = {
            "COMMENTS": self.comments,
            "KEYWORDS": self.keywords,
            "SUBJECT": self.subject,
            "TITLE": self.title
        }

        # Get the function from switcher dictionary
        func = switcher.get(argument, lambda: "Invalid property")

        # Execute the function
        func(val)

    #--------------------------------------------------------------------------#
    # Function: processTables                                                  #
    #                                                                          #
    # Description:                                                             #
    # Processes all of the tables in the document looking.                     #
    #--------------------------------------------------------------------------#
    def processTables(self):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'processTables'

        #----------------------------------------------------------------------#
        # Define global progress bar variable:                                 #
        #----------------------------------------------------------------------#
        global ps

        #----------------------------------------------------------------------#
        # Enter a loop to process each table in the document:                  #
        #----------------------------------------------------------------------#
        numTables = len(self.document.tables)
        iTable = 0
        for table in self.document.tables:
            #------------------------------------------------------------------#
            # Process the table and update any SQL data:                       #
            #------------------------------------------------------------------#
            iTable = iTable + 1
            self.dataTable(table)

            #------------------------------------------------------------------#
            # Update the progress bar:                                         #
            #------------------------------------------------------------------#
            pc = 1.0 / numTables
            ps.update(pc)
            ps.set_description('Updating data in table ' + str(iTable))
            ps.refresh()

    #--------------------------------------------------------------------------#
    # Function: dataTable                                                      #
    #                                                                          #
    # Description:                                                             #
    # Processes the current table looking for defined header keywords defining #
    # the content of the table.                                                #
    #--------------------------------------------------------------------------#
    def dataTable(self, table):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'dataTable'

        #----------------------------------------------------------------------#
        # Get the table data source:                                           #
        #----------------------------------------------------------------------#
        bDelRow = False
        self.currentTable = table
        self.currentRow = self.currentTable.rows[0]
        txtQuery = self.currentRow.cells[0].text

        #----------------------------------------------------------------------#
        # Check for non-ascii characters. Can't be a keyword string in that    #
        # case:                                                                #
        #----------------------------------------------------------------------#
        try:
            txtQuery.decode('ascii')
        except:
            errorHandler(errProc, errorCode.nonASCIICharacter, txtQuery)
        else:
            #------------------------------------------------------------------#
            # Check for the content type. Check if an image anchor:            #
            #------------------------------------------------------------------#
            if (txtQuery[:5] == 'IMAGE'):
                #--------------------------------------------------------------#
                # Insert the image:                                            #
                #--------------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableInsertImage(txtQuery[6:])

            #------------------------------------------------------------------#
            # Check if a base query record table:                              #
            #------------------------------------------------------------------#
            elif (txtQuery[:7] == 'SQLPROP'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.docProperty(txtQuery[8:])

            #------------------------------------------------------------------#
            # Check if an append table rows query:                             #
            #------------------------------------------------------------------#
            elif (txtQuery[:6] == 'SQLROW'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableAddRows(txtQuery[7:])

            #------------------------------------------------------------------#
            # Check if a static table query:                                   #
            #------------------------------------------------------------------#
            elif (txtQuery[:9] == 'SQLSTATIC'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                self.tableStaticFields(txtQuery[10:])

            #------------------------------------------------------------------#
            # Check if an append table rows query:                             #
            #------------------------------------------------------------------#
            elif (txtQuery[:10] == 'SQLVERHIST'):
                #--------------------------------------------------------------#
                # Must be a valid query. Insert the data:                      #
                #--------------------------------------------------------------#
                self.remove_row(self.currentTable, self.currentRow)
                self.currentRow = self.currentTable.rows[0]
                query = cgSQL.sql[cgSQL.sqlCode.VERHIST]
                self.tableAddRows(query)
            else:
                #--------------------------------------------------------------#
                # Ignore other content:                                        #
                #--------------------------------------------------------------#
                pass

    #--------------------------------------------------------------------------#
    # Function: docProperty                                                    #
    #                                                                          #
    # Description:                                                             #
    # Sets the document property based on the query result.                    #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def docProperty(self, query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'docProperty'

        #----------------------------------------------------------------------#
        # Execute the query and get the data. Static table can only handle     #
        # one record:                                                          #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        data = c.fetchone()
        if (not data is None):
            #------------------------------------------------------------------#
            # Try to set the document property if the field exists:            #
            #------------------------------------------------------------------#
            try:
                prop = data.keys()
                self.setProperty(prop[0], data[0])
            except:
                pass

    #--------------------------------------------------------------------------#
    # Function: tableAddRows                                                   #
    #                                                                          #
    # Description:                                                             #
    # Adds one table row for each row in the cursor.                           #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string to get the data for.              #
    #--------------------------------------------------------------------------#
    def tableAddRows(self, query):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableAddRows'

        #----------------------------------------------------------------------#
        # Execute the query and get the data:                                  #
        #----------------------------------------------------------------------#
        bHasData = False
        c = self.getQueryData(query)
        data = c.fetchall()

        #----------------------------------------------------------------------#
        # Iterate the data to get the row count:                               #
        #----------------------------------------------------------------------#
        self.rowCount = 0
        for row in data:
            bHasData = True
            self.rowCount = self.rowCount + 1

        #----------------------------------------------------------------------#
        # There is no data so delete the entire table and the paragraph:       #
        #----------------------------------------------------------------------#
        if (not bHasData):
            p1 = self.getTableParagraph(self.currentTable)
            self.remove_table(self.currentTable)
            self.remove_paragraph(p1)
        else:
            #------------------------------------------------------------------#
            # Re-query to refresh the cursor:                                  #
            #------------------------------------------------------------------#
            c = self.getQueryData(query)
            data = c.fetchall()

            #------------------------------------------------------------------#
            # Get the column attribute row:                                    #
            #------------------------------------------------------------------#
            rowAttr = self.currentTable.rows[1]
            cellsAttr = rowAttr.cells
            para = cellsAttr[0].paragraphs[0]
            styleAttr = para.style

            #------------------------------------------------------------------#
            # Process each row of returned data:                               #
            #------------------------------------------------------------------#
            for row in data:
                #--------------------------------------------------------------#
                # Add a new row to the table and enter a loop to process each  #
                # cell:                                                        #
                #--------------------------------------------------------------#
                rowNew = self.currentTable.add_row()
    #            rowNew.height_rule = WD_ROW_HEIGHT.EXACTLY
                rowNew.height_rule = 2
                rowNew.height = shared.Cm(1.2)
    #            cellsNew = self.currentTable.add_row().cells
                cellsNew = rowNew.cells
                for i in range(0, len(cellsAttr)):
                    #----------------------------------------------------------#
                    # Replace the field placeholders in the cell text:         #
                    #----------------------------------------------------------#
                    s = cellsAttr[i].text
                    for fld in row.keys():
                        try:
                            s = s.replace('@@' + fld.upper() + '@@', str(row[fld]))
                        except:
                            pass
                    para = cellsNew[i].paragraphs[0]
                    para.text = ''
                    run = para.add_run(s)
                    para.style = styleAttr

            #------------------------------------------------------------------#
            # Clean up the table by deleting the query and attribute rows:     #
            #------------------------------------------------------------------#
            self.remove_row(self.currentTable, rowAttr)

    #--------------------------------------------------------------------------#
    # Function: tableInsertImage                                               #
    #                                                                          #
    # Description:                                                             #
    # Inserts an image into the table cell.                                    #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # sImage                The image filename.                                #
    #--------------------------------------------------------------------------#
    def tableInsertImage(self, sImage):
        #----------------------------------------------------------------------#
        # Define the procedure name and trap any programming errors:           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableInsertImage'

        #----------------------------------------------------------------------#
        # Get the image file:                                                  #
        #----------------------------------------------------------------------#
        for fld in self.dataRow.keys():
            sImage = sImage.replace('@@' + fld.upper() + '@@', str(self.dataRow[fld]))

        #----------------------------------------------------------------------#
        # Check if the image exists:                                           #
        #----------------------------------------------------------------------#
        sImage = self.inputDir + '/' + sImage
        if not os.path.exists(sImage):
            errorHandler(self.errProc, errorCode.fileNotExist, sImage)

        #----------------------------------------------------------------------#
        # Insert the image:                                                    #
        #----------------------------------------------------------------------#
        para = self.currentRow.cells[0].paragraphs[0]
        style = para.style
        para.text = ''
        run = para.add_run()
        run.add_picture(sImage)
        para.style = style

    #--------------------------------------------------------------------------#
    # Function: tableStaticFields                                              #
    #                                                                          #
    # Description:                                                             #
    # Populates a table with placeholder data without altering the table       #
    # structure.                                                               #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # query                 The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def tableStaticFields(self, query):
        #----------------------------------------------------------------------#
        # Define the procedure name:                                           #
        #----------------------------------------------------------------------#
        self.errProc = 'tableStaticFields'

        #----------------------------------------------------------------------#
        # Execute the query and get the data. Static table can only handle     #
        # one record:                                                          #
        #----------------------------------------------------------------------#
        c = self.getQueryData(query)
        data = c.fetchone()
        if (data is None):
            #------------------------------------------------------------------#
            # No data so delete the table and the paragraph:                   #
            #------------------------------------------------------------------#
            p1 = self.getTableParagraph(self.currentTable)
            self.remove_table(self.currentTable)
            self.remove_paragraph(p1)
        else:
            #------------------------------------------------------------------#
            # Process each field in the data and all cells in the table:       #
            #------------------------------------------------------------------#
            for fld in data.keys():
                #--------------------------------------------------------------#
                # Ignore non-ascii characters:                                 #
                #--------------------------------------------------------------#
                try:
                    str(data[fld]).decode('ascii')
                    self.srTable(self.currentTable, '@@' + fld.upper() + '@@', str(data[fld]))
                except:
                    errorHandler(self.errProc, errorCode.nonASCIIField, query, fld)

    #--------------------------------------------------------------------------#
    # Function: getQueryData                                                   #
    #                                                                          #
    # Description:                                                             #
    # Queries the database with the specified SQL string.                      #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryData(self, txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of programming errors:             #
        #----------------------------------------------------------------------#
        self.errProc = 'getQueryData'

        #----------------------------------------------------------------------#
        # Get any parameters from the query string:                            #
        #----------------------------------------------------------------------#
        q = self.getQueryParameters(txtQuery)

        #----------------------------------------------------------------------#
        # Execute the query:                                                   #
        #----------------------------------------------------------------------#
        try:
            c = self.conn.cursor()
            c.execute(q.query, q.parms)
        except:
            errorHandler(self.errProc, errorCode.cannotQuery, q.query, q.parms)
        return c

    #--------------------------------------------------------------------------#
    # Function: getQueryParameters                                             #
    #                                                                          #
    # Description:                                                             #
    # Gets the parameter values from the parameter names in the query string.  #
    #--------------------------------------------------------------------------#
    # Calling parameters:                                                      #
    #                                                                          #
    # txtQuery              The query string for the table data.               #
    #--------------------------------------------------------------------------#
    def getQueryParameters(self, txtQuery):
        #----------------------------------------------------------------------#
        # Define the procedure name in case of any programming errors:         #
        #----------------------------------------------------------------------#
        errProc = 'getQueryParameters'

        #----------------------------------------------------------------------#
        # Declare the named tuple for the returned data:                       #
        #----------------------------------------------------------------------#
        Query = collections.namedtuple('Query', ['query', 'parms'])

        #----------------------------------------------------------------------#
        # Declare local parameter list to return:                              #
        #----------------------------------------------------------------------#
        parms = []

        #----------------------------------------------------------------------#
        # Enter a loop to process all of the parameter placeholders in the     #
        # query:                                                               #
        #----------------------------------------------------------------------#
        while (txtQuery.find('@@') >= 0):
            #------------------------------------------------------------------#
            # Search for any parameter placeholders in the query string:       #
            #------------------------------------------------------------------#
            iTagBegin = txtQuery.find('@@')
            iTagEnd = txtQuery.find('@@', iTagBegin + 2)
            if (iTagBegin == -1):
                #--------------------------------------------------------------#
                # Nothing to do. Return null:                                  #
                #--------------------------------------------------------------#
                return;

            elif (iTagBegin > iTagEnd):
                errorHandler(self.errProc, errorCode.noEndPlaceholder, txtQuery)
            else:
                #--------------------------------------------------------------#
                # Get the parameter text between the BEGIN and END tags and    #
                # replace the parameter with a question mark. Add the          #
                # parameter to the list:                                       #
                #--------------------------------------------------------------#
                sParameterName = txtQuery[iTagBegin + 2:iTagEnd]
                txtQuery = txtQuery[:iTagBegin] + '?' + txtQuery[iTagEnd + 2:]
                parms.append(sParameterName)

        #----------------------------------------------------------------------#
        # Replace the parameters with the existing property values:            #
        #----------------------------------------------------------------------#
        for i in range(len(parms)):
            if (parms[i].upper() == 'FILTER'):
                parms[i] = self.filter

        #----------------------------------------------------------------------#
        # Return the parameter value list:                                     #
        #----------------------------------------------------------------------#
        q = Query(query=txtQuery, parms=parms)
        return q;

    def srParagraph(self, paragraph, txtSearch, txtReplace):
        errProc = 'srParagraph'

        #----------------------------------------------------------------------#
        # Update the text in the paragraph by replacing the search text:       #
        #----------------------------------------------------------------------#
        s = paragraph.text
        if (s.find(txtSearch) >= 0):
            style = paragraph.style
#            txtReplace = txtReplace.replace('\"','')
            s = s.replace(txtSearch, txtReplace)
            paragraph.text = ''
            run = paragraph.add_run(s)
            paragraph.style = style

    def srDocument(self, document, txtSearch, txtReplace):
        errProc = 'srDocument'

        for paragraph in self.document.paragraphs:
            self.srParagraph(paragraph, txtSearch, txtReplace)

    def srHeader(self, document, txtSearch, txtReplace):
        errProc = 'srHeader'

        for section in self.document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self.srParagraph(paragraph, txtSearch, txtReplace)

    def srTable(self, table, txtSearch, txtReplace):
        errProc = 'srTable'

        for row in self.currentTable.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.srParagraph(paragraph, txtSearch, txtReplace)

    def remove_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def remove_row(self, t, r):
        errProc = 'remove_row'
        tbl = t._tbl
        tr = r._tr
        tbl.remove(tr)

    def remove_table(self, t):
        errProc = 'remove_table'
        tbl = self.currentTable._tbl
        for row in self.currentTable.rows:
            tr = row._tr
            tbl.remove(tr)
#        tbl = t._tbl
#        self.document.tables.remove(tbl)

    #--------------------------------------------------------------------------#
    # Return a newly created paragraph, inserted directly before this          #
    # item (Table, etc.):                                                      #
    #--------------------------------------------------------------------------#
#    def insert_paragraph_before(self, item, text, style=None):
#        p = CT_P.add_p_before(item._element)
#        p2 = Paragraph(p, item._parent)
#        p2.text = text
#        p2.style = style
#        return p2

    #--------------------------------------------------------------------------#
    # Yield each paragraph and table child within *parent*, in document order. #
    # Each returned value is an instance of either Table or Paragraph. *parent*#
    # would most commonly be a reference to a main Document object, but        #
    # also works for a _Cell object, which itself can contain paragraphs and   #
    # tables:                                                                  #
    #--------------------------------------------------------------------------#
    def iter_block_items(self, parent):
        if isinstance(parent, type(self.document)):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def getTableParagraph(self, t):
        nodePrev = self.document.paragraphs[0]

        for node in self.iter_block_items(self.document):
            if isinstance(node, Table):
                if (node._tbl == t._tbl):
                    return nodePrev
            else:
                nodePrev = node

#------------------------------------------------------------------------------#
# Call the main function:                                                      #
#------------------------------------------------------------------------------#
main()
