#------------------------------------------------------------------------------#
#            Copyright 2018 Rieckermann Engineering Operations                 #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates a word report based on the module class.                   #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Rev By               Date        CC        Note                              #
# 1.0 David Paspa      05-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
from docx import Document
import openpyxl

#------------------------------------------------------------------------------#
# Create a new document:                                                       #
#------------------------------------------------------------------------------#
document = Document()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2)
table = document.add_table(rows=2, cols=2)
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph.style = 'Normal'
document.save('./test.odt')

#------------------------------------------------------------------------------#
# Open the workbook:                                                           #
#------------------------------------------------------------------------------#
wb = openpyxl.load_workbook('io.xlsx')
#wb.get_sheet_names()
#['Sheet1', 'Sheet2', 'Sheet3']
sheet = wb['Sheet1']
#type(sheet) <class 'openpyxl.worksheet.worksheet.Worksheet'>
print(sheet.title)
#'Sheet3'
#anotherSheet = wb.active
#anotherSheet
#<Worksheet "Sheet1">
