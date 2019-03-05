#------------------------------------------------------------------------------#
#                   Copyright 2018 complianceSHORTCUTS.com                     #
#------------------------------------------------------------------------------#
# Description:                                                                 #
# This file creates document from a sqlite database by searching and replacing #
# the table field names as @@name@@ data placeholders within the document.     #
#------------------------------------------------------------------------------#
# Revision history:                                                            #
# Ver By               Date        CC        Note                              #
# 2   David Paspa      27-Jul-2018 NA        Split from document generator     #
#                                            into separate module. Added       #
#                                            multiple prefix handling.         #
# 1   David Paspa      13-Apr-2018 NA        Initial design.                   #
#------------------------------------------------------------------------------#
# Import the python libraries:                                                 #
#------------------------------------------------------------------------------#
import argparse
import re
from enum import Enum
import os.path
import sys
from collections import defaultdict
import operator
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font
from openpyxl.styles.borders import Border, Side
from openpyxl.utils import get_column_letter
#from openpyxl.workbook import Workbook
from openpyxl import Workbook
#from openpyxl.worksheet.properties import WorksheetProperties
import traceback
from tqdm import trange
from fdgDoc import gDoc

import logging
logging.basicConfig(filename='fdg.log',level=logging.DEBUG)

#------------------------------------------------------------------------------#
# Define constants:                                                            #
#------------------------------------------------------------------------------#
COL_WIDTH_LABEL = 10
COL_WIDTH_X = 3
FILL_REQ = 'C5D9F1'
FILL_TEST = 'EBF1DE'
ROW_TITLE = 2

#------------------------------------------------------------------------------#
# Declare the application title and calling arguments help:                    #
#------------------------------------------------------------------------------#
appTitle = 'TSB Generator'
appVersion = '1'
parser = argparse.ArgumentParser(description='Generates an Traceability Switchboard')
parser.add_argument('-i','--inputFile', help='Input document file to cross reference', required=True)
parser.add_argument('-o','--outputFile', help='The output TSB file name', required=True)
parser.add_argument('-p','--listPrefix', help='The reference number prefixes', required=True)
parser.add_argument('-r','--listReferencesPrefix', help='The reference document number prefixes', required=True)
parser.add_argument('-t','--tag', help='The tag number prefix for the TSB', required=True)
parser.add_argument('-l','--left', help='Put the references on the left, otherwise right', required=True)

#------------------------------------------------------------------------------#
# Declare global variables:                                                    #
#------------------------------------------------------------------------------#
colTag = 0
refData = None
refLeft = True
refTag = ''
rpDefs = None
ps = None
wb = None
ws = None

#------------------------------------------------------------------------------#
# Get the arguments and store in local variables:                              #
#------------------------------------------------------------------------------#
args = vars(parser.parse_args())
inputFile = args['inputFile']
outputFile = args['outputFile']
prefixes = args['listPrefix']
references = args['listReferencesPrefix']
refTag = args['tag']

#------------------------------------------------------------------------------#
# Get the left or right orientation:                                           #
#------------------------------------------------------------------------------#
def str_to_bool(s):
    if (s.upper() == 'TRUE'):
        return True
    elif (s.upper() == 'FALSE'):
        return False
    else:
        raise ValueError("Cannot covert {} to a bool".format(s))

refLeft = str_to_bool(args['left'])

#------------------------------------------------------------------------------#
# Declare the error handling global variables and procedure:                   #
#------------------------------------------------------------------------------#
def errorHandler(errProc, eCode, *args):
    #--------------------------------------------------------------------------#
    # Declare global variables:                                                #
    #--------------------------------------------------------------------------#
    global ps

    #--------------------------------------------------------------------------#
    # Get the application specific error message and output the error:         #
    #--------------------------------------------------------------------------#
    sMsg = errorMessage[eCode]

    #--------------------------------------------------------------------------#
    # Replace any error parameters:                                            #
    #--------------------------------------------------------------------------#
    i = 1
    for arg in args:
        sMsg = sMsg.replace('@' + str(i), str(arg))
        i = i + 1

    #--------------------------------------------------------------------------#
    # Close the progress bar:                                                  #
    #--------------------------------------------------------------------------#
    ps.close()

    #--------------------------------------------------------------------------#
    # Output the error message and end:                                        #
    #--------------------------------------------------------------------------#
    print('\r\n')
    print('ERROR ' + str(eCode) + ' in Procedure ' + "'" + errProc + "'" + '\r\n' + '\r\n' + sMsg)
    print('\r\n')
#    print(traceback.format_exception(*sys.exc_info()))
    sys.exit()

#------------------------------------------------------------------------------#
# Enumerate the error numbers and map the error messages:                      #
#------------------------------------------------------------------------------#
class errorCode(Enum):
    fileNotExist                       = -1
    noNamedRange                       = -2
    noPrefixFound                      = -3
    pathNotExist                       = -4

errorMessage = {
    errorCode.fileNotExist             : 'Document file @1 does not exist.',
    errorCode.noNamedRange             : 'Named range @1 does not exist.',
    errorCode.noPrefixFound            : 'No reference numbers found with the specified prefix @1.',
    errorCode.pathNotExist             : 'Path @1 does not exist.'
}

#------------------------------------------------------------------------------#
# Function main                                                                #
#                                                                              #
# Description:                                                                 #
# The main entry point for the program.                                        #
#------------------------------------------------------------------------------#
def main():
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = main.__name__

    #--------------------------------------------------------------------------#
    # Use global tag number variable:                                          #
    #--------------------------------------------------------------------------#
    global colTag
    global refData
    global refTag
    global ps
    global wb
    global ws

    #--------------------------------------------------------------------------#
    # Create a progress bar:                                                   #
    #--------------------------------------------------------------------------#
    ps = trange(1, desc='Initialise', leave=False)

    #--------------------------------------------------------------------------#
    # Get the calling arguments:                                               #
    #--------------------------------------------------------------------------#
    listPrefix = prefixes.split(',')
    listRefPrefix = references.split(',')

    #--------------------------------------------------------------------------#
    # Get the input document:                                                  #
    #--------------------------------------------------------------------------#
    d = gDoc(inputFile, inputFile)

    #--------------------------------------------------------------------------#
    # Get the TSB workbook name and check it exists:                           #
    #--------------------------------------------------------------------------#
    wbName = args['outputFile']
    if not os.path.exists(wbName):
        #----------------------------------------------------------------------#
        # Create a new workbook as it doesn't exist yet:                       #
        #----------------------------------------------------------------------#
        wb = openpyxl.Workbook()
        ws = wb.active

        #----------------------------------------------------------------------#
        # Format the worksheet as a new TSB:                                   #
        #----------------------------------------------------------------------#
        tsbFormat(d.fileInputBaseName, listPrefix[0])
        wb.save(wbName)
    else:
        #----------------------------------------------------------------------#
        # Open the existing workbook:                                          #
        #----------------------------------------------------------------------#
        try:
            wb = openpyxl.load_workbook(wbName, data_only=True)
            ws = wb.active
        except:
            errorHandler(errProc, errorCode.cannotOpenWorkbook, wbName)

    #--------------------------------------------------------------------------#
    # Get the Tag column number:                                               #
    #--------------------------------------------------------------------------#
    try:
        sNamedRange = 'Tag'
        colTag = ws[list(wb.defined_names[sNamedRange].destinations)[0][1]].col_idx
    except:
        errorHandler(errProc, errorCode.noNamedRange, sNamedRange)

    #--------------------------------------------------------------------------#
    # Get the reference prefix names and column numbers:                       #
    #--------------------------------------------------------------------------#
    refData = defaultdict(list)
    rpDefs = getRefPrefixDefinitions(d, listRefPrefix)

    #--------------------------------------------------------------------------#
    # Generate the TSB:                                                        #
    #--------------------------------------------------------------------------#
    bFoundPrefix = getRefData(d, listPrefix, listRefPrefix)

    #--------------------------------------------------------------------------#
    # Generate the TSB if any reference data found:                            #
    #--------------------------------------------------------------------------#
    if (bFoundPrefix):
        rpData = sorted(refData.items(), key=lambda kv: kv[0])
        for r in rpData:
            tsbXRef(refTag + str(r[0]), r[1])

    #--------------------------------------------------------------------------#
    # Highlight any unreferenced requirements or tests:                        #
    #--------------------------------------------------------------------------#
    showUnreferenced(True)
    showUnreferenced(False)

    #--------------------------------------------------------------------------#
    # Define the border cell formats:                                          #
    #--------------------------------------------------------------------------#
    tsbBorder = Border(left=Side(style='hair'), right=Side(style='hair'),
                       top=Side(style='hair'), bottom=Side(style='hair'))
    tsbBorderBoxLeft = Border(left=Side(style='hair'), right=Side(style='none'),
                              top=Side(style='hair'), bottom=Side(style='hair'))
    tsbBorderBoxMiddle = Border(left=Side(style='none'), right=Side(style='none'),
                                top=Side(style='hair'), bottom=Side(style='hair'))
    tsbBorderBoxRight = Border(left=Side(style='none'), right=Side(style='hair'),
                               top=Side(style='hair'), bottom=Side(style='hair'))

    #--------------------------------------------------------------------------#
    # Set the borders for all cells:                                           #
    #--------------------------------------------------------------------------#
    colMax = ws.max_column
    rowMax = ws.max_row
    iRow = ROW_TITLE
    for iRow in range (2, rowMax + 1):
        for iCol in range (1, colMax + 1):
            ws.cell(row=iRow, column=iCol).border = tsbBorder

    #--------------------------------------------------------------------------#
    # Set the left title row borders:                                          #
    #--------------------------------------------------------------------------#
    ws.cell(row=1, column=1).border = tsbBorderBoxLeft
    ws.cell(row=1, column=colTag - 1).border = tsbBorderBoxRight
    for iCol in range (2, colTag - 1):
        ws.cell(row=1, column=iCol).border = tsbBorderBoxMiddle

    #--------------------------------------------------------------------------#
    # Set the right title row borders:                                         #
    #--------------------------------------------------------------------------#
    ws.cell(row=1, column=colTag + 1).border = tsbBorderBoxLeft
    ws.cell(row=1, column=colMax).border = tsbBorderBoxRight
    for iCol in range (colTag + 2, colMax):
        ws.cell(row=1, column=iCol).border = tsbBorderBoxMiddle

    #--------------------------------------------------------------------------#
    # Report completion regardless of error:                                   #
    #--------------------------------------------------------------------------#
    wb.save(wbName)
    if (bFoundPrefix):
        ps.set_description('TSB generated successfully.')
        ps.refresh()
    else:
        errorHandler(errProc, errorCode.noPrefixFound, sPrefix)
    ps.close()

#------------------------------------------------------------------------------#
# Function: getRefData                                                         #
#                                                                              #
# Description:                                                                 #
# Gets the reference number cross-reference data from the document.            #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# d                     The generic document object.                           #
# listPrefix            The prefix list to look through.                       #
# listRefPrefix         The reference document prefix list to look through.    #
#------------------------------------------------------------------------------#
def getRefData(d, listPrefix, listRefPrefix):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = getRefData.__name__

    #--------------------------------------------------------------------------#
    # Use global tag number variable:                                          #
    #--------------------------------------------------------------------------#
    global colTag
    global refData
    global refTag
    global rpDefs
    global wb
    global ws

    #--------------------------------------------------------------------------#
    # Loop through all the tables in the document table collection:            #
    #--------------------------------------------------------------------------#
    refDataRaw = defaultdict(list)
    tagIdxMax = 0
    refIdxMax = 0
    bFoundPrefix = False
    iTable = 0
    numTables = len(d.document.tables)
    pn = re.compile('[0-9]')
    for table in d.document.tables:
        iRow = 0
        iTable = iTable + 1
        for row in table.rows:
            numRows = len(table.rows)
            iRow = iRow + 1
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    #----------------------------------------------------------#
                    # Check if a valid tag number:                             #
                    #----------------------------------------------------------#
                    s = paragraph.text
                    for p in listPrefix:
                        if (len(s) >= 1 + len(p) and s[:len(p)] == p):
                            mn = pn.match(s[len(p):])
                            if (not mn is None):
                                tagIdx = int(s[len(p):])
                                if (tagIdx > tagIdxMax):
                                    tagIdxMax = tagIdx
                                tagPrefix = p
                                bFoundPrefix = True
                                for cell in row.cells:
                                    #------------------------------------------#
                                    # Add any valid reference document number  #
                                    # cross references:                        #
                                    #------------------------------------------#
                                    rs = cell.text
                                    rpList = rs.split('\n')
                                    for rp in rpList:
                                        for rpd in rpDefs:
                                            if (len(rp) >= 1 + len(rpd[0]) and rp[:len(rpd[0])] == rpd[0]):
                                                mn = pn.match(rp[len(rpd[0]):])
                                                if (not mn is None):
                                                    refIdx = int(rp[len(rpd[0]):])
                                                    if (refTag == listPrefix[0]):
                                                        refDataRaw[tagIdx].append(rp)

                                                    elif (refTag == rpd[0]):
                                                        if (refIdx > refIdxMax):
                                                            refIdxMax = refIdx
                                                        refDataRaw[refIdx].append(tagPrefix + str(tagIdx))

            #------------------------------------------------------------------#
            # Update the progress bar:                                         #
            #------------------------------------------------------------------#
            pc = 1.0 / (numTables * numRows)
            ps.update(pc)
            ps.set_description('Processing table ' + str(iTable) + ' row ' + str(iRow))
            ps.refresh()

    #--------------------------------------------------------------------------#
    # Arrange the data with all reference numbers to indicate missing:         #
    #--------------------------------------------------------------------------#
    if (refIdxMax > 0):
        maxNum = refIdxMax
    else:
        maxNum = tagIdxMax
    for x in range(1, maxNum + 1):
        if (len(refDataRaw[x]) == 0):
            refData[x].append('')
        else:
            for r in refDataRaw[x]:
                refData[x].append(r)

    #--------------------------------------------------------------------------#
    # Return whether any prefix was found or not:                              #
    #--------------------------------------------------------------------------#
    return bFoundPrefix

#------------------------------------------------------------------------------#
# Function: tsbXRef                                                            #
#                                                                              #
# Description:                                                                 #
# Adds a cross reference to the TSB.                                           #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# s                     The reference tag number.                              #
# rpList                The reference document tag number list from the doc.   #
#------------------------------------------------------------------------------#
def tsbXRef(s, rpList):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = tsbXRef.__name__

    #--------------------------------------------------------------------------#
    # Constants:                                                               #
    #--------------------------------------------------------------------------#
    FILL_TAG = 'BFBFBFBF'

    #--------------------------------------------------------------------------#
    # Use global worksheet variable:                                           #
    #--------------------------------------------------------------------------#
    global colTag
    global rpDefs
    global wb
    global ws

    #--------------------------------------------------------------------------#
    # Search for the reference in the tag list:                                #
    #--------------------------------------------------------------------------#
    rh = None
    rowRef = 1
    for c in ws[get_column_letter(colTag)]:
        if (c.value == s):
            #------------------------------------------------------------------#
            # The reference number row has been found:                         #
            #------------------------------------------------------------------#
            rowRef = c.row
            break

    #--------------------------------------------------------------------------#
    # Check if the tag number was not found:                                   #
    #--------------------------------------------------------------------------#
    if (rowRef == 1):
        rowRef = len(ws[get_column_letter(colTag)]) + 1
        c = ws.cell(row=rowRef, column=colTag)
        c.alignment = Alignment(horizontal='center', vertical='top', text_rotation=0, wrap_text=True, shrink_to_fit=False, indent=0)
        c.fill = PatternFill(fill_type='solid', start_color=FILL_TAG, end_color=FILL_TAG)
        c.value = s

    #--------------------------------------------------------------------------#
    # Process all of the prefixes:                                             #
    #--------------------------------------------------------------------------#
    pn = re.compile('[0-9]')
    for rpd in rpDefs:
        #----------------------------------------------------------------------#
        # Don't add references for the root tag:                               #
        #----------------------------------------------------------------------#
        if (refTag != rpd[0]):
            #------------------------------------------------------------------#
            # Get the existing reference tag data as a list:                   #
            #------------------------------------------------------------------#
            r = ws.cell(row=rowRef, column=rpd[2])
            rh = r.value
            if (rh is None):
                rh = ''
            rhList = rh.split('\n')

            #------------------------------------------------------------------#
            # Delete the old references which have been removed:               #
            #------------------------------------------------------------------#
            hDelete = set(rhList).difference(rpList)
            rhList = [x for x in rhList if x not in hDelete]

            #------------------------------------------------------------------#
            # Process all of the references:                                   #
            #------------------------------------------------------------------#
            for rp in rpList:
                #--------------------------------------------------------------#
                # Add the new references to the TSB if a matching prefix:      #
                #--------------------------------------------------------------#
                if (len(rp) >= 1 + len(rpd[0]) and rp[:len(rpd[0])] == rpd[0]):
                    mn = pn.match(rp[len(rpd[0]):])
                    if (not mn is None):
                        rpoList = []
                        rpoList.append(rp)
                        hMissing = set(rpoList).difference(rhList)
                        rhList.extend(hMissing)

                #--------------------------------------------------------------#
                # Sort the updated reference number list:                      #
                #--------------------------------------------------------------#
                convert = lambda text: int(text) if text.isdigit() else text
                alphanum_key = lambda key: [convert(c) for c in re.split('([0-9]+)', key)]
                rhnn = sorted(rhList, key = alphanum_key)

                #--------------------------------------------------------------#
                # Update the reference document tags and exit:                 #
                #--------------------------------------------------------------#
                rhnx = ''
                for h in rhnn:
                    if (len(rhnx) > 0):
                        rhnx = rhnx + '\n'
                    rhnx = rhnx + h
                r.value = rhnx
                r.alignment = Alignment(horizontal='left', vertical='top', text_rotation=0, wrap_text=True, shrink_to_fit=False, indent=0)

#------------------------------------------------------------------------------#
# Function: tsbFormat                                                          #
#                                                                              #
# Description:                                                                 #
# Format a new TSB workbook.                                                   #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# sTitle                The TSB Title.                                         #
# sPrefix               The valid prefix of the document.                      #
#------------------------------------------------------------------------------#
def tsbFormat(sTitle, sPrefix):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = tsbFormat.__name__

    #--------------------------------------------------------------------------#
    # Use global worksheet constants and variables:                            #
    #--------------------------------------------------------------------------#
    global COL_WIDTH_LABEL
    global COL_WIDTH_X
    global FILL_REQ
    global FILL_TEST
    global ROW_TITLE
    global colTag
    global wb
    global ws

    #--------------------------------------------------------------------------#
    # Create the named ranges:                                                 #
    #--------------------------------------------------------------------------#
    colTag = 4
    createNamedRanges(colTag)

    #--------------------------------------------------------------------------#
    # Format the cells:                                                        #
    #--------------------------------------------------------------------------#
    c = ws.cell(row=ROW_TITLE - 1, column=3)
    c.value = 'Requirements'
    c.alignment=Alignment(horizontal='right', vertical='top', text_rotation=0, wrap_text=False, shrink_to_fit=False, indent=0)
    c.font = Font(bold=True)
    c = ws.cell(row=ROW_TITLE - 1, column=5)
    c.value = 'Tests'
    c.alignment=Alignment(horizontal='left', vertical='top', text_rotation=0, wrap_text=False, shrink_to_fit=False, indent=0)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag - 3)
#    c.value = 'Reference'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_REQ, end_color=FILL_REQ)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag - 2)
    c.value = 'Justification'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_REQ, end_color=FILL_REQ)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag - 1)
    c.value = 'Unreferenced'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_REQ, end_color=FILL_REQ)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag)
    c.value = sTitle
    c.alignment = Alignment(horizontal='center', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag + 1)
    c.value = 'Untested'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_TEST, end_color=FILL_TEST)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag + 2)
    c.value = 'Justification'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_TEST, end_color=FILL_TEST)
    c.font = Font(bold=True)

    c = ws.cell(row=ROW_TITLE, column=colTag + 3)
#    c.value = 'Test'
    c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90, wrap_text=False, shrink_to_fit=False, indent=0)
    c.fill = PatternFill(fill_type='solid', start_color=FILL_TEST, end_color=FILL_TEST)
    c.font = Font(bold=True)

    #--------------------------------------------------------------------------#
    # Set the row heights and column widths:                                   #
    #--------------------------------------------------------------------------#
#    ws.row_dimensions[3].height = 1
    ws.column_dimensions[get_column_letter(colTag - 3)].width = COL_WIDTH_LABEL
    ws.column_dimensions[get_column_letter(colTag - 2)].width = COL_WIDTH_X
    ws.column_dimensions[get_column_letter(colTag - 1)].width = COL_WIDTH_X
    ws.column_dimensions[get_column_letter(colTag)].width = COL_WIDTH_LABEL
    ws.column_dimensions[get_column_letter(colTag + 1)].width = COL_WIDTH_X
    ws.column_dimensions[get_column_letter(colTag + 2)].width = COL_WIDTH_X
    ws.column_dimensions[get_column_letter(colTag + 3)].width = COL_WIDTH_LABEL

    #--------------------------------------------------------------------------#
    # Set the page properties:                                                 #
    #--------------------------------------------------------------------------#
    ws.oddHeader.center.text = 'Traceability Switchboard ' + sTitle
    ws.oddHeader.center.size = 10
    ws.oddFooter.center.text = "Page &[Page] of &N"
    ws.oddFooter.center.size = 8
#    ws.oddHeader.right.font = "Tahoma,Bold"
#    ws.oddHeader.right.color = "CC3366"
#    ws.print_title_cols = 'A:B' # the first two cols
    ws.print_title_rows = '1:2'
    ws.print_options.horizontalCentered = True
#    ws.print_options.verticalCentered = True
#    ws.sheet_view.showGridLines = True

#------------------------------------------------------------------------------#
# Function: createNamedRanges                                                  #
#                                                                              #
# Description:                                                                 #
# Creates the named ranges in the workbook.                                    #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# colTag                The tag column number.                                 #
#------------------------------------------------------------------------------#
def createNamedRanges(colTag):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = createNamedRanges.__name__

    #--------------------------------------------------------------------------#
    # Use global constants:                                                    #
    #--------------------------------------------------------------------------#
    global COL_WIDTH_LABEL
    global COL_WIDTH_X
    global ROW_TITLE

    #--------------------------------------------------------------------------#
    # Create the named ranges for the columns:                                 #
    #--------------------------------------------------------------------------#
    sNamedRange = 'Reference'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag - 3) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag - 3)].width = COL_WIDTH_LABEL
    sNamedRange = 'JustificationReference'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag - 2) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag - 2)].width = COL_WIDTH_X
    sNamedRange = 'Unreferenced'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag - 1) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag - 1)].width = COL_WIDTH_X
    sNamedRange = 'Tag'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag)].width = COL_WIDTH_LABEL
    sNamedRange = 'Untested'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag + 1) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag + 1)].width = COL_WIDTH_X
    sNamedRange = 'JustificationTest'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag + 2) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag + 2)].width = COL_WIDTH_X
    sNamedRange = 'Test'
    wb.create_named_range(sNamedRange, ws, '$' + get_column_letter(colTag + 3) + '$' + str(ROW_TITLE))
    ws.column_dimensions[get_column_letter(colTag + 3)].width = COL_WIDTH_LABEL

#------------------------------------------------------------------------------#
# Function: deleteNamedRanges                                                  #
#------------------------------------------------------------------------------#
def deleteNamedRanges():
    del wb.defined_names['Reference']
    del wb.defined_names['JustificationReference']
    del wb.defined_names['Unreferenced']
    del wb.defined_names['Tag']
    del wb.defined_names['Untested']
    del wb.defined_names['JustificationTest']
    del wb.defined_names['Test']

#------------------------------------------------------------------------------#
# Function: getRefPrefixDefinitions                                            #
#                                                                              #
# Description:                                                                 #
# Gets the reference prefix name from the Prefix table in the document and the #
# column number from the TSB spreadsheet.                                      #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# d                     The generic document object.                           #
# listRefPrefix         The reference prefix list to look through.             #
#------------------------------------------------------------------------------#
# return                The tuple list of prefix definitions.                  #
#------------------------------------------------------------------------------#
def getRefPrefixDefinitions(d, listRefPrefix):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = getRefPrefixDefinitions.__name__

    #--------------------------------------------------------------------------#
    # Use the global reference document prefix definitions:                    #
    #--------------------------------------------------------------------------#
    global rpDefs

    #--------------------------------------------------------------------------#
    # Loop through all the tables in the document table collection:            #
    #--------------------------------------------------------------------------#
    rpDefs = []
    for rp in listRefPrefix:
        #----------------------------------------------------------------------#
        # Get the prefix name and column number and add the reference prefix   #
        # definition tuple to the list:                                        #
        #----------------------------------------------------------------------#
        rpName = getRefPrefixName(d, rp)
        rpCol = 0
        if (refTag != rp):
            rpCol = getRefPrefixColumn(rpName)
        rpDefs.append(tuple((rp, rpName, rpCol)))

    #--------------------------------------------------------------------------#
    # Return the list of tuples:                                               #
    #--------------------------------------------------------------------------#
    return rpDefs

#------------------------------------------------------------------------------#
# Function: updateRefPrefixDefinitions                                         #
#                                                                              #
# Description:                                                                 #
# Updates the reference prefix column numbers in the immutable tuple.          #
#------------------------------------------------------------------------------#
# return                Updated tuple list of prefix definitions.              #
#------------------------------------------------------------------------------#
def updateRefPrefixDefinitions():
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = updateRefPrefixDefinitions.__name__

    #--------------------------------------------------------------------------#
    # Use the global reference document prefix definitions:                    #
    #--------------------------------------------------------------------------#
    global rpDefs

    #--------------------------------------------------------------------------#
    # Loop through all the reference prefixes:                                 #
    #--------------------------------------------------------------------------#
    rpnDefs = []
    for rp in rpDefs:
        #----------------------------------------------------------------------#
        # Get the prefix name and column number and add the reference prefix   #
        # definition tuple to the list:                                        #
        #----------------------------------------------------------------------#
        rpPrefix = rp[0]
        rpName = rp[1]
        rpCol = 0
        if (refTag != rp):
            rpCol = getRefPrefixColumn(rp[1])
        rpnDefs.append(tuple((rpPrefix, rpName, rpCol)))

    #--------------------------------------------------------------------------#
    # Return the list of tuples:                                               #
    #--------------------------------------------------------------------------#
    return rpnDefs

#------------------------------------------------------------------------------#
# Function: getRefPrefixName                                                   #
#                                                                              #
# Description:                                                                 #
# Gets the reference prefix name from the Prefix table in the document.        #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# d                     The generic document object.                           #
# sRefPrefix            The reference prefix to find.                          #
#------------------------------------------------------------------------------#
# return                The reference document prefix name.                    #
#------------------------------------------------------------------------------#
def getRefPrefixName(d, sRefPrefix):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = getRefPrefixName.__name__

    #--------------------------------------------------------------------------#
    # Loop through all the tables in the document table collection:            #
    #--------------------------------------------------------------------------#
    for table in d.document.tables:
        for row in table.rows:
            iCell = 0
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    #----------------------------------------------------------#
                    # Check if a prefix table entry:                           #
                    #----------------------------------------------------------#
                    s = paragraph.text
                    if (s == sRefPrefix):
                        iCell = iCell + 1
                        return row.cells[iCell].paragraphs[0].text

    #--------------------------------------------------------------------------#
    # Return the document name if not found:                                   #
    #--------------------------------------------------------------------------#
    return d.fileInputBaseName

#------------------------------------------------------------------------------#
# Function: getRefPrefixColumn                                                 #
#                                                                              #
# Description:                                                                 #
# Gets the reference prefix column number from the TSB. Adds a new column if   #
# it doesn't already exist.                                                    #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# sRefName              The reference document name to look for.               #
#------------------------------------------------------------------------------#
# return                The reference document prefix column number.           #
#------------------------------------------------------------------------------#
def getRefPrefixColumn(sRefName):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = getRefPrefixColumn.__name__

    #--------------------------------------------------------------------------#
    # Use global constants:                                                    #
    #--------------------------------------------------------------------------#
    global COL_WIDTH_LABEL
    global FILL_REQ
    global FILL_TEST
    global ROW_TITLE

    #--------------------------------------------------------------------------#
    # Use global worksheet variable:                                           #
    #--------------------------------------------------------------------------#
    global colTag
    global refLeft
    global rpDefs
    global ws

    #--------------------------------------------------------------------------#
    # Get the reference column named range to orientate the data:              #
    #--------------------------------------------------------------------------#
    colReference = colTag - 3
    colTest = colTag + 3

    #--------------------------------------------------------------------------#
    # Start at the reference column:                                           #
    #--------------------------------------------------------------------------#
    foundCol = False
    if (refLeft):
        rpCol = colReference
    else:
        rpCol = colTest
    c = ws.cell(row=ROW_TITLE, column=rpCol)
    s = c.value
    if (s is None):
        s = ''

    #--------------------------------------------------------------------------#
    # Check all reference columns for the prefix:                              #
    #--------------------------------------------------------------------------#
    while(rpCol > 0 and len(s) > 0):
        c = ws.cell(row=ROW_TITLE, column=rpCol)
        s = c.value
        if (s == sRefName):
            foundCol = True
            break
        else:
            if (refLeft):
                rpCol = rpCol - 1
            else:
                rpCol = rpCol + 1

    #--------------------------------------------------------------------------#
    # Check if the reference document does not exist:                          #
    #--------------------------------------------------------------------------#
    if (not foundCol and rpCol == 0 and rpCol != colReference and rpCol != colTest):
        #----------------------------------------------------------------------#
        # Add a column for the new reference document:                         #
        #----------------------------------------------------------------------#
        rpCol = 1
        ws.insert_cols(rpCol, 1)

        #----------------------------------------------------------------------#
        # Move the named ranges to the right:                                  #
        #----------------------------------------------------------------------#
        deleteNamedRanges()
        colTag = colTag + 1
        createNamedRanges(colTag)

        #----------------------------------------------------------------------#
        # Get the reference prefix tuple definitions with the updated column   #
        # numbers:                                                             #
        #----------------------------------------------------------------------#
        rpnDefs = []
        rpnDefs = updateRefPrefixDefinitions()
        rpDefs = rpnDefs

    #--------------------------------------------------------------------------#
    # Add the reference document name:                                         #
    #--------------------------------------------------------------------------#
    if (not foundCol):
        c = ws.cell(row=ROW_TITLE, column=rpCol)
        c.value = sRefName
        c.alignment = Alignment(horizontal='left', vertical='bottom', text_rotation=90)
        c.font = Font(bold=True)
        ws.column_dimensions[get_column_letter(rpCol)].width = COL_WIDTH_LABEL
        if (refLeft):
            c.fill = PatternFill(fill_type='solid', start_color=FILL_REQ, end_color=FILL_REQ)
        else:
            c.fill = PatternFill(fill_type='solid', start_color=FILL_TEST, end_color=FILL_TEST)

    #--------------------------------------------------------------------------#
    # Return the tuple list:                                                   #
    #--------------------------------------------------------------------------#
    return rpCol

#------------------------------------------------------------------------------#
# Function: showUnreferenced                                                   #
#                                                                              #
# Description:                                                                 #
# Shows any unreferenced items.                                                #
#------------------------------------------------------------------------------#
# Calling Parameters:                                                          #
# showLeft              True if unreferenced requirements, otherwise tests.    #
#------------------------------------------------------------------------------#
def showUnreferenced(showLeft):
    #--------------------------------------------------------------------------#
    # Define the procedure name:                                               #
    #--------------------------------------------------------------------------#
    errProc = showUnreferenced.__name__

    #--------------------------------------------------------------------------#
    # Use global worksheet variable:                                           #
    #--------------------------------------------------------------------------#
    global colTag
    global ws

    #--------------------------------------------------------------------------#
    # Get the preset column positions:                                         #
    #--------------------------------------------------------------------------#
    colReference = colTag - 3
    colTest = colTag + 3
    if (showLeft):
        colUnreferenced = colTag - 1
    else:
        colUnreferenced = colTag + 1

    #--------------------------------------------------------------------------#
    # Process all of the Tag numbers in the TSB list:                          #
    #--------------------------------------------------------------------------#
    colMax = ws.max_column
    rowMax = ws.max_row
    rowRef = ROW_TITLE + 1
    while(rowRef <= rowMax):
        #----------------------------------------------------------------------#
        # Start at the reference or test column:                               #
        #----------------------------------------------------------------------#
        foundX = False
        if (showLeft):
            rpCol = colReference
        else:
            rpCol = colTest

        #----------------------------------------------------------------------#
        # Check all reference columns:                                         #
        #----------------------------------------------------------------------#
        while(rpCol > 0 and rpCol < colMax + 1):
            c = ws.cell(row=rowRef, column=rpCol)
            s = c.value
            if (s is None):
                s = ''
            if (len(s) > 0):
                foundX = True
                break
            else:
                if (showLeft):
                    rpCol = rpCol - 1
                else:
                    rpCol = rpCol + 1

        #----------------------------------------------------------------------#
        # Check if there are references:                                       #
        #----------------------------------------------------------------------#
        c = ws.cell(row=rowRef, column=colUnreferenced)
        if (foundX):
            c.value = ''
        else:
            #------------------------------------------------------------------#
            # No references so highlight the row:                              #
            #------------------------------------------------------------------#
            c.alignment = Alignment(horizontal='center', vertical='center')
            c.value = 'X'

        #----------------------------------------------------------------------#
        # Check the next row:                                                  #
        #----------------------------------------------------------------------#
        rowRef = rowRef + 1

#------------------------------------------------------------------------------#
# Call the main function:                                                      #
#------------------------------------------------------------------------------#
main()
